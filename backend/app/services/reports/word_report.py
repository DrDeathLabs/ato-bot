"""Word findings narrative report."""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from sqlalchemy import select

from app.core.config import get_settings
from app.core.database import AsyncSessionLocal
from app.models.orm import Assessment, ControlFinding, Project

settings = get_settings()

STATUS_COLORS = {
    "compliant": RGBColor(0x00, 0x70, 0x00),
    "partially_compliant": RGBColor(0xFF, 0x8C, 0x00),
    "non_compliant": RGBColor(0xCC, 0x00, 0x00),
    "not_applicable": RGBColor(0x80, 0x80, 0x80),
    "not_reviewed": RGBColor(0x00, 0x00, 0x00),
}

# NIST SP 800-53A Rev 5 formal determination language for SAR output
NIST_DETERMINATION = {
    "compliant": "Satisfied",
    "partially_compliant": "Other Than Satisfied",
    "non_compliant": "Other Than Satisfied",
    "not_applicable": "Not Applicable",
    "not_reviewed": "Not Reviewed",
}
INTERNAL_LABELS = {
    "compliant": "Compliant",
    "partially_compliant": "Partially Compliant",
    "non_compliant": "Non-Compliant",
    "not_applicable": "Not Applicable",
    "not_reviewed": "Not Reviewed",
}


async def generate_word(assessment_id: int) -> str:
    async with AsyncSessionLocal() as db:
        result = await db.execute(select(Assessment).where(Assessment.id == assessment_id))
        assessment = result.scalar_one()
        proj_result = await db.execute(select(Project).where(Project.id == assessment.project_id))
        project = proj_result.scalar_one()
        findings_result = await db.execute(
            select(ControlFinding)
            .where(ControlFinding.assessment_id == assessment_id)
            .order_by(ControlFinding.control_family, ControlFinding.control_id)
        )
        findings = findings_result.scalars().all()

    doc = Document()

    # Title
    title = doc.add_heading("Security Assessment Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"System: {project.name}").runs[0].bold = True
    doc.add_paragraph(f"Baseline: {project.impact_baseline.upper()}")
    doc.add_paragraph(f"Date: {assessment.completed_at.strftime('%B %d, %Y') if assessment.completed_at else 'In Progress'}")
    doc.add_paragraph(f"Assessment ID: {assessment_id}")
    doc.add_page_break()

    # Executive Summary
    doc.add_heading("Executive Summary", 1)
    compliant = sum(1 for f in findings if f.status == "compliant")
    partial = sum(1 for f in findings if f.status == "partially_compliant")
    non = sum(1 for f in findings if f.status == "non_compliant")
    na = sum(1 for f in findings if f.status == "not_applicable")
    total = len(findings)

    summary_text = (
        f"This report presents the results of an automated security assessment of {project.name} "
        f"against the NIST SP 800-53 Rev 5 {project.impact_baseline.upper()} baseline "
        f"({total} controls assessed).\n\n"
        f"Results: {compliant} Compliant ({compliant/total*100:.0f}%) | "
        f"{partial} Partially Compliant | {non} Non-Compliant | {na} Not Applicable"
    )
    doc.add_paragraph(summary_text)
    doc.add_page_break()

    # Group by family
    families = sorted({f.control_family for f in findings})
    for family in families:
        doc.add_heading(f"Control Family: {family}", 1)
        fam_findings = [f for f in findings if f.control_family == family]

        for finding in fam_findings:
            # Control heading
            p = doc.add_heading(f"{finding.control_id} — {finding.control_title}", 2)

            # Obs 8: formal NIST 800-53A determination + internal label in parentheses
            # e.g. "Determination: Other Than Satisfied (Partially Compliant)"
            nist_label = NIST_DETERMINATION.get(finding.status, "Not Reviewed")
            internal_label = INTERNAL_LABELS.get(finding.status, finding.status.replace("_", " ").title())
            status_para = doc.add_paragraph()
            status_run = status_para.add_run(f"Determination: {nist_label}")
            status_run.bold = True
            color = STATUS_COLORS.get(finding.status, RGBColor(0, 0, 0))
            status_run.font.color.rgb = color
            if nist_label != internal_label:
                qualifier = status_para.add_run(f"  ({internal_label})")
                qualifier.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
                qualifier.font.size = Pt(9)

            # Obs 8: synthesized narrative notice — analyst and AO need to know
            if getattr(finding, "synthesized_narrative", False):
                notice_para = doc.add_paragraph()
                notice_run = notice_para.add_run(
                    "⚠ Note: Implementation statement was auto-synthesized from gap analysis. "
                    "A full Stage 3 narrative pass was not performed. "
                    "Re-run this control to produce a complete SSP narrative before final submission."
                )
                notice_run.italic = True
                notice_run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
                notice_run.font.size = Pt(9)

            # Implementation statement
            if finding.implementation_statement:
                doc.add_heading("Implementation Statement", 4)
                doc.add_paragraph(finding.implementation_statement)

            # Gaps
            if finding.gaps:
                doc.add_heading("Identified Gaps", 4)
                for gap in finding.gaps:
                    doc.add_paragraph(gap, style="List Bullet")

            # Remediation
            if finding.remediation_plan and finding.status in ("partially_compliant", "non_compliant"):
                doc.add_heading("Remediation Plan", 4)
                doc.add_paragraph(finding.remediation_plan)

            # AI dissent note — recorded second-opinion; code verdict remains authoritative
            if getattr(finding, "llm_challenge_note", None):
                doc.add_heading("AI Assessor Dissent Note", 4)
                dissent_para = doc.add_paragraph(finding.llm_challenge_note)
                dissent_para.runs[0].italic = True
                dissent_para.runs[0].font.color.rgb = RGBColor(0x55, 0x00, 0x88)
                doc.add_paragraph(
                    "Note: The automated verdict is authoritative. "
                    "This note is a recorded second opinion for human assessor review.",
                ).runs[0].font.size = Pt(8)

            # Evidence citations
            if finding.evidence_citations:
                doc.add_heading("Evidence Citations", 4)
                for cite in finding.evidence_citations[:3]:
                    if isinstance(cite, dict):
                        p = doc.add_paragraph()
                        p.add_run(f"{cite.get('source', '')}: ").bold = True
                        p.add_run(cite.get('quote', ''))

    output_path = Path(settings.output_dir) / f"assessment_{assessment_id}_findings.docx"
    os.makedirs(settings.output_dir, exist_ok=True)
    doc.save(str(output_path))
    return str(output_path)
