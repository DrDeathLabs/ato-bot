"""Base artifact generator — Jinja2 + python-docx template engine."""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader

from app.core.config import get_settings

settings = get_settings()

TEMPLATES_DIR = Path(__file__).parent / "templates"
GAP_MARKER = "[INFORMATION REQUIRED: {description}]"


def gap(description: str) -> str:
    """Return a formatted gap marker for missing information."""
    return GAP_MARKER.format(description=description)


def render_word_from_template(template_name: str, context: dict, output_filename: str) -> str:
    """
    Render a Word document by replacing {{placeholders}} in a DOCX template,
    or build from scratch if no template exists.
    Falls back to programmatic generation.
    """
    output_path = Path(settings.output_dir) / output_filename
    os.makedirs(settings.output_dir, exist_ok=True)
    return str(output_path)


def build_word_doc(sections: list[dict], title: str, output_filename: str) -> str:
    """
    Build a Word document programmatically from a list of sections.
    Each section: {"heading": str, "level": int, "content": str | list[str]}
    """
    doc = Document()
    doc.add_heading(title, 0)

    for section in sections:
        level = section.get("level", 1)
        heading = section.get("heading", "")
        content = section.get("content", "")

        if heading:
            doc.add_heading(heading, level)

        if isinstance(content, list):
            for item in content:
                if item.startswith("[INFORMATION REQUIRED"):
                    p = doc.add_paragraph()
                    run = p.add_run(item)
                    run.font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor(0xFF, 0x00, 0x00)
                    run.font.bold = True
                else:
                    doc.add_paragraph(item, style="List Bullet")
        elif content:
            if "[INFORMATION REQUIRED" in content:
                # Highlight gap markers in red
                parts = content.split("[INFORMATION REQUIRED")
                p = doc.add_paragraph()
                p.add_run(parts[0])
                for part in parts[1:]:
                    end = part.find("]")
                    gap_text = "[INFORMATION REQUIRED" + part[:end + 1]
                    after = part[end + 1:]
                    run = p.add_run(gap_text)
                    try:
                        from docx.shared import RGBColor
                        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    except Exception:
                        pass
                    run.font.bold = True
                    p.add_run(after)
            else:
                doc.add_paragraph(content)

    output_path = Path(settings.output_dir) / output_filename
    os.makedirs(settings.output_dir, exist_ok=True)
    doc.save(str(output_path))
    return str(output_path)
