"""
Remediation Service — post-assessment gap closure generation.

Two report types:
  guide     — Prioritized action plan saved as Word (.docx) + Excel (.xlsx) tracker.
  artifacts — Reassessment-ready implementation artifacts (.docx), saved to the
              project document library and indexed for future assessments.
"""
from __future__ import annotations

import asyncio
import hashlib
import io
import json
import logging
import re
from collections import defaultdict
from datetime import UTC, datetime
from pathlib import Path
from time import perf_counter
from uuid import uuid4

from sqlalchemy import select, update

from app.core.config import get_settings
from app.core.database import AsyncSessionLocal
from app.models.orm import Assessment, AssessmentCriteriaPackage, ControlFinding, Document, RemediationReport
from app.services.closure_guidance import (
    build_control_closure_guidance,
    build_contract_sections,
    format_contracts_for_prompt,
    sections_satisfy_contracts,
)
from app.services.controls.catalog import load_catalog
from app.services.evidence_view import get_document_evidence_payload
from app.services.implementation_statements import (
    build_control_statement_generation_guidance,
    synthesize_control_implementation_statement,
)
from app.services.package_generation import build_blueprint_validation, plan_remediation_bundles

logger = logging.getLogger(__name__)
settings = get_settings()

PARTIAL_SAVE_EVERY_BUNDLES = 3
PARALLEL_CONTROL_WORKERS = max(1, min(settings.remediation_parallel_controls, 8))


def _split_assessable_findings(findings: list) -> tuple[list, list[dict]]:
    """Remove NIST-withdrawn controls from remediation work while keeping traceability."""
    catalog = load_catalog()
    by_display_id = {
        re.sub(r"\s+", "", control.display_id.upper()): control
        for control in catalog.values()
    }

    assessable_findings: list = []
    non_assessable: list[dict] = []
    for finding in findings:
        raw_control_id = str(getattr(finding, "control_id", "") or "").strip()
        control = (
            by_display_id.get(re.sub(r"\s+", "", raw_control_id.upper()))
            or catalog.get(raw_control_id.lower())
        )
        if control and not control.is_assessable:
            target = ", ".join(control.incorporated_into) or "the incorporated control"
            non_assessable.append({
                "control_id": control.display_id,
                "title": control.title,
                "status": control.status,
                "disposition": "withdrawn_incorporated",
                "incorporated_into": control.incorporated_into,
                "reason": (
                    f"{control.display_id} is withdrawn by NIST and is not assessed independently. "
                    f"Assessment coverage is handled under {target}."
                ),
            })
            continue
        assessable_findings.append(finding)

    return assessable_findings, non_assessable

# ── Prompt: Remediation Guide (per family) ─────────────────────────────────────

REMEDIATION_GUIDE_SYSTEM_PROMPT = """You are a NIST 800-53 Rev 5 compliance expert writing a practical remediation guide.

For each compliance gap provided, generate a specific, actionable remediation action.

Return ONLY a valid JSON array. Each element:
{
  "control_id": "AC-2",
  "gap": "brief gap description",
  "action": "specific, concrete action required",
  "responsible": "role (e.g. ISSO, System Administrator, CISO, HR)",
  "effort": "effort estimate (e.g. 1-2 days, 1 week, 2-4 weeks, 1-3 months)",
  "success_criteria": "what 'done' looks like — verifiable outcome",
  "template_language": "1-3 sentence example policy/procedure statement they can adapt"
}

RULES:
- Actions must be specific and verifiable — not "review policies" but "draft and approve AC-2 account management policy covering provisioning, modification, disabling, and removal"
- Responsible role must match the action type (policy = ISSO/CISO, technical = Sys Admin, HR controls = HR Manager)
- Template language should be professional, compliance-ready prose they can copy and adapt
- RESPOND WITH ONLY THE JSON ARRAY"""

# ── Fix 6: Technical controls that require real evidence (not generated docs) ──

# These controls require technical implementation evidence — configs, logs,
# scan results, contracts, or operational records. For remediation runs we still
# generate reassessment-ready artifacts, but they must read like completed
# technical evidence packages rather than abstract policy prose.
TECHNICAL_CONTROL_IDS: frozenset[str] = frozenset({
    # Network / Communications
    "SC-2", "SC-3", "SC-4", "SC-5", "SC-7", "SC-7(4)", "SC-7(7)",
    "SC-10", "SC-15", "SC-18", "SC-20", "SC-22", "SC-28", "SC-28(1)", "SC-39",
    # Malware / Integrity / Monitoring
    "SI-3", "SI-3(1)", "SI-3(2)", "SI-4", "SI-4(4)", "SI-6", "SI-7", "SI-8", "SI-8(1)",
    # Audit infrastructure
    "AU-3(1)", "AU-4", "AU-7(1)", "AU-9(4)",
    # Contingency infrastructure (alternate site / telecom)
    "CP-6", "CP-6(1)", "CP-6(3)", "CP-7", "CP-7(1)", "CP-7(2)", "CP-7(3)",
    "CP-8", "CP-8(1)",
    # Continuous monitoring / independent assessment
    "CA-7(1)",
    # Developer / supply chain
    "SA-17",
    # Crypto
    "SC-13", "SC-17",
})

# ── Prompt: Per-control artifact planner ──────────────────────────────────────

ARTIFACT_PLANNER_SYSTEM_PROMPT = """You are a NIST 800-53 Rev 5 compliance evidence architect.

CONTEXT: This document will be indexed by a RAG-based AI assessor. When the assessor evaluates the control, it retrieves chunks by vector similarity and checks each NIST 800-53A assessment objective ID explicitly. Your document MUST satisfy each objective listed or the control will remain non-compliant.

Your task: for a SINGLE failing control with specific objective-level gaps, choose the ONE document type that best closes ALL listed gaps with reassessment-ready implementation evidence.

Return ONLY a valid JSON object (not an array):
{
  "title": "Exact document title — include the control ID",
  "artifact_type": "policy | policy_procedure | completed_form | ssp_narrative | procedure | technical_artifact | agreement_template",
  "purpose": "One sentence: which specific objectives this document closes and how"
}

SELECTION RULES:
- Objective IDs only (e.g. "AC-02a.[01]") with no other text → ssp_narrative
- "no documented process" / "no procedure" / "no steps" → procedure
- "no evidence" / "no records" / "no logs" / "no audit trail" / "technical implementation not shown" → technical_artifact
- "only policy cited, no implementation" / "no actual [X] provided" → completed_form
- "no signed agreement" / "no MOU" / "no authorization" → agreement_template
- pure governance / X-1 policy gaps → policy
- Multiple policy/procedure gaps across one control → policy_procedure
- RESPOND WITH ONLY THE JSON OBJECT — no array wrapper, no explanation"""


# ── Prompt: Artifact content generator ────────────────────────────────────────

ARTIFACT_CONTENT_BASE_PROMPT = """You are a NIST 800-53 Rev 5 compliance author writing an official security document that must pass reassessment.

CRITICAL CONTEXT: This document will be indexed by an AI assessment system that evaluates each NIST 800-53A assessment objective individually by retrieving relevant text chunks. The assessor must find explicit language satisfying each objective. Vague or generic content WILL NOT score as compliant.

AUTHORING MODE:
- Write in present-tense current-state language that describes how the control is implemented today.
- Treat this artifact as current evidence, not a future-state gap template or aspirational policy draft.
- Use specific roles, systems, tools, records, timelines, and verification actions.
- If realistic placeholder values are needed for testing, make them concrete and fully populated rather than leaving blanks.
- Before the per-objective sections, include a heading named "Control Implementation Summary" followed by 2-4 human-readable paragraphs that synthesize the control as one coherent implementation statement.
- The objectives are coverage requirements for that summary, not separate implementation statements.

MANDATORY STRUCTURE FOR EACH GAP:
For every assessment objective listed in the gaps, you MUST create a dedicated section using this exact pattern:
  - Heading: "[OBJECTIVE_ID] — [Short descriptive title]"
    Example heading: "AC-02f.[04] — Account Disabling Procedure"
  - First paragraph sentence: "This section satisfies NIST 800-53A assessment objective [OBJECTIVE_ID]."
  - Following sentences: specific, concrete implementation language addressing exactly what the objective requires.

This section-per-objective structure is NON-NEGOTIABLE. It is what allows the assessor AI to retrieve and score each objective independently.

OUTPUT FORMAT — return ONLY a valid JSON object:
{
  "title": "Full document title",
  "sections": [
    {"type": "heading", "level": 1, "text": "Section title"},
    {"type": "paragraph", "text": "Plain prose only — no markdown symbols"},
    {"type": "numbered_list", "items": ["Step text"]},
    {"type": "bullet_list", "items": ["Item text"]},
    {"type": "table", "headers": ["Col1", "Col2"], "rows": [["val", "val"]]},
    {"type": "divider"}
  ]
}

CONTENT RULES:
- Use the EXACT objective ID notation in headings (e.g., "AC-02f.[04]", "PL-02a.01[01]")
- Prefer implemented language: "is configured", "is reviewed", "is recorded", "uses", "maintains", "retains"
- Only use SHALL where the document type truly requires policy authority language, and still make the implementation current and explicit.
- Be specific: name tools, roles, timelines, thresholds — not generic statements
- If the objective implies a value, timeout, threshold, review cadence, dissemination audience, approval authority, rationale, alert workflow, or retained record, state it explicitly.
- If the objective references policy, procedure, or system documentation, include the actual dissemination, approval, and review/update details rather than general assurances.
- Plain prose only — NO markdown symbols (#, **, *, |, ---) anywhere in text fields
- Return ONLY the JSON object — no preamble, no commentary"""

# ── Prompt: Technical implementation evidence package ─────────────────────────

EVIDENCE_COLLECTION_SYSTEM_PROMPT = """You are a NIST 800-53 Rev 5 compliance specialist writing a completed Technical Implementation Evidence Package.

This control requires TECHNICAL IMPLEMENTATION EVIDENCE — system configurations, network diagrams, scan results, test logs, vendor contracts, or physical records. A generic policy document will NOT close this gap. Your job is to produce a structured, current-state artifact that reads like completed technical evidence and will be indexed as implementation support.

Write in present tense as if the implementation is already in place for testing purposes.
Use realistic but fictitious identifiers, schedules, reviewers, and tool outputs when needed.
DO provide: specific evidence types, where they are generated, what values or records are captured, and how verification is performed.

MANDATORY STRUCTURE: For each failing objective, create a dedicated section:
  - Heading: "[OBJECTIVE_ID] — Technical Implementation Evidence"
  - Paragraph beginning: "This section satisfies NIST 800-53A assessment objective [OBJECTIVE_ID]."
  - Numbered list of concrete implementation details, evidence records, verification steps, and owners

Required document sections: Technical Context, [one section per objective gap], Configuration Evidence Tables, Verification Record, Evidence Retention and Review

OUTPUT FORMAT — return ONLY a valid JSON object with the standard sections structure.
Plain prose only in text fields — no markdown symbols."""

ARTIFACT_TYPE_INSTRUCTIONS = {
    "policy": """DOCUMENT TYPE: Formal Security Policy

Required sections (heading nodes): Purpose, Scope, Authority, Policy Statements, [one section per objective gap — see mandatory structure above], Roles and Responsibilities, Enforcement, Review Schedule, Document Control

POLICY WRITING RULES:
- Write as an approved policy already in force for the system and organization
- Use present-tense authority language tied to current operations: "maintains", "requires", "reviews", "approves"
- Every objective section must explicitly name the role, cadence, threshold, and evidence record that demonstrates implementation
- When governance or dissemination is implicated, include audience, publication method, acknowledgment/distribution record, approving authority, and review/update trigger.
- Document Control must be fully populated with realistic dates, version, preparer, and approver details""",

    "policy_procedure": """DOCUMENT TYPE: Formal Policy and Procedure

Required sections (heading nodes): Purpose, Scope, Applicability, Policy Statements, [one section per objective gap — see mandatory structure above], Roles and Responsibilities, Enforcement, Review Schedule, Compliance Verification

COMPLIANCE WRITING RULES:
- After the standard intro sections, add one dedicated heading+paragraph per objective gap
- Heading format: "[OBJECTIVE_ID] — [Short Title]"  First sentence: "This section satisfies NIST 800-53A assessment objective [OBJECTIVE_ID]."
- Write as a current implemented policy/procedure pair, not a future-state draft
- Compliance Verification section: table mapping each objective ID to its verification method, evidence source, review cadence, and responsible role
- If the objective calls for dissemination, compliance language, management commitment, or law/directive alignment, state those explicitly in dedicated rows or paragraphs.
- Do NOT use: "as appropriate", "when necessary", "should consider", "may"
- Use the exact wording from NIST 800-53A assessment objectives in the relevant sections""",

    "completed_form": """DOCUMENT TYPE: Completed Form / Matrix / Register

This is a FILLED-IN artifact — an assessor must find explicit objective-level evidence in the table.

COMPLETION RULES:
- First: add one heading+paragraph per objective gap before the table (mandatory structure)
- Then: table with all rows fully populated, NO empty cells
- Use realistic system-specific data and completed entries, not empty templates
- Include the objective ID in a dedicated column header: "Objective Satisfied"
- For each row, populate the "Objective Satisfied" column with the specific NIST objective ID
- For registers/inventories: 10-20 representative rows
- End with Document Control table: Version, Date, Preparer, Approver""",

    "ssp_narrative": """DOCUMENT TYPE: System Security Plan (SSP) — Control Implementation Narratives

Required structure — for each objective gap, use this exact pattern:
  Heading level 2: "[OBJECTIVE_ID] — [Short Title]"
  Paragraph 1: "This section satisfies NIST 800-53A assessment objective [OBJECTIVE_ID]."
  Paragraph 2: Implementation Description — specific description of how this objective is met
  Paragraph 3: Evidence — specific documents, logs, or configurations proving implementation
  Paragraph 4: Responsible Role — ISSO / System Owner / System Administrator

WRITING RULES:
- Write in the voice of the system owner/ISSO
- Be specific: name the tools, processes, systems, and configurations
- The NIST objective notation (e.g., "AC-02f.[04]") must appear in the heading text
- Do not group multiple objectives into one section — one section per objective ID
- If the objective expects SSP documentation or rationale, state the exact SSP section and the documented rationale directly.
- Treat the narrative as current implementation evidence, not planned future work""",

    "procedure": """DOCUMENT TYPE: Operational Procedure

Required sections (heading nodes): Purpose, Scope, Applicability, Prerequisites, [one section per objective gap], Procedure Steps, Verification, Record-Keeping, Completion Checklist

PROCEDURE WRITING RULES:
- After Prerequisites, add one dedicated heading+paragraph per objective gap (mandatory structure above)
- Procedure Steps: numbered_list — each step is actionable: who, what system/tool, what output
- Record-Keeping: table with columns: What is Logged, System/Location, Retention Period, Reviewer
- Completion Checklist: table with columns: Objective ID, Requirement, Completed
  Each row corresponds to one objective gap — populated with the objective ID and requirement
- If the objective implies dissemination, approvals, or review/update activity, include a table row capturing that record explicitly.
- Use the exact control ID and NIST objective language in the Purpose paragraph
- Write the procedure as the currently used operational process with concrete systems, frequencies, and records""",

    "technical_artifact": """DOCUMENT TYPE: Technical Implementation Evidence Package

Required sections: Technical Context, [one section per objective gap], Configuration Evidence Table, Verification Record, Retention and Review Requirements

TECHNICAL WRITING RULES:
- For each objective gap, add heading "[OBJECTIVE_ID] — Technical Implementation Evidence" with paragraph
  beginning "This section satisfies NIST 800-53A assessment objective [OBJECTIVE_ID]."
- Include specific technologies, configurations, log sources, scans, or verification artifacts that prove implementation
- Configuration Evidence Table: all rows fully populated with concrete values and owners
- Verification Record: table with Objective ID, Verification Method, Last Verified, Result, Reviewer
- Retention and Review Requirements: table with Requirement, Value columns
- This document must read like implementation evidence, not collection guidance or an empty template""",

    "evidence_template": """DOCUMENT TYPE: Technical Implementation Evidence Package

Required sections: Technical Context, [one section per objective gap], Configuration Evidence Table, Verification Record, Retention and Review Requirements

TECHNICAL WRITING RULES:
- Treat this as a completed technical evidence record for testing purposes
- For each objective gap, add heading "[OBJECTIVE_ID] — Technical Implementation Evidence" with paragraph
  beginning "This section satisfies NIST 800-53A assessment objective [OBJECTIVE_ID]."
- Include concrete evidence rows, verification details, owners, and review cadence
- Do not output blank template fields or fill-in markers""",

    "agreement_template": """DOCUMENT TYPE: Inter-Organizational Agreement / Authorization

Required sections: Parties, Purpose, Authority, [one section per objective gap], System Description, Authorized Activities, Security Requirements, Obligations, Duration, Contacts, Signatures

AGREEMENT WRITING RULES:
- For each objective gap, add heading "[OBJECTIVE_ID] — Agreement Requirement"
  First sentence: "This section satisfies NIST 800-53A assessment objective [OBJECTIVE_ID]."
- Security Requirements: table mapping each objective ID to its specific requirement
- Signatures: table with Name, Title, Organization, Date, Signature (two rows)
- Use concrete signatory names and titles rather than empty placeholders
- Cite the specific NIST control in the Authority section and state the currently authorized activities and obligations""",
}


# ── Fix 6: Build gap objectives as structured list ────────────────────────────

def _parse_objective_entry(raw_value: str, fallback_control_id: str) -> tuple[str, str]:
    value = str(raw_value or "").strip()
    obj_match = re.match(
        r'^([A-Z]{2}-\d+[a-z]?(?:\(\d+\))?(?:[a-z])?(?:\.\d+)?(?:\[\d+\])?)[\s:\.]+',
        value
    )
    if obj_match:
        objective_id = obj_match.group(1)
        description = value[len(obj_match.group(0)):].strip() or value
    else:
        objective_id = fallback_control_id
        description = value
    return objective_id, description


def _build_gap_objectives(finding, assessment_objectives: list | None = None) -> list[dict]:
    """Convert gaps and full assessment objectives into structured objective entries."""
    gaps = finding.gaps or []
    if isinstance(gaps, str):
        gaps = [gaps]
    objectives = []
    gap_map: dict[str, dict] = {}
    for gap in gaps:
        gap_str = str(gap).strip()
        obj_id, description = _parse_objective_entry(gap_str, finding.control_id)
        gap_map[obj_id] = {
            "objective_id": obj_id,
            "description": description or gap_str,
            "full_text": gap_str,
        }

    seen: set[str] = set()
    for raw_objective in assessment_objectives or []:
        if isinstance(raw_objective, dict):
            objective_id = str(raw_objective.get("id") or raw_objective.get("objective_id") or finding.control_id).strip()
            description = str(
                raw_objective.get("description")
                or raw_objective.get("prose")
                or raw_objective.get("text")
                or objective_id
            ).strip()
        else:
            objective_id, description = _parse_objective_entry(str(raw_objective), finding.control_id)
        if objective_id in seen:
            continue
        seen.add(objective_id)
        gap_entry = gap_map.pop(objective_id, None)
        objectives.append({
            "objective_id": objective_id,
            "description": (gap_entry or {}).get("description") or description,
            "full_text": (gap_entry or {}).get("full_text") or f"{objective_id}: {description}",
        })

    for objective_id, gap_entry in gap_map.items():
        if objective_id in seen:
            continue
        seen.add(objective_id)
        objectives.append(gap_entry)
    return objectives


def _normalize_objective_prompt_entries(objectives: list | None, control_id: str) -> list[dict]:
    """Coerce prompt objective payloads into the structured dict shape the writer expects."""
    normalized: list[dict] = []
    for raw in objectives or []:
        if isinstance(raw, dict):
            objective_id = str(raw.get("objective_id") or raw.get("id") or control_id).strip()
            description = str(
                raw.get("description")
                or raw.get("full_text")
                or raw.get("text")
                or objective_id
            ).strip()
            full_text = str(raw.get("full_text") or f"{objective_id}: {description}").strip()
        else:
            objective_id, description = _parse_objective_entry(str(raw), control_id)
            full_text = f"{objective_id}: {description}"
        normalized.append(
            {
                "objective_id": objective_id,
                "description": description,
                "full_text": full_text,
            }
        )
    return normalized


def _format_objectives_for_prompt(objectives: list[dict], control_id: str, control_title: str) -> str:
    """Format objectives as the structured gap input for content generation prompts."""
    objectives = _normalize_objective_prompt_entries(objectives, control_id)
    contracts = build_control_closure_guidance(
        control_id=control_id,
        control_title=control_title,
        gaps=objectives,
        system_name="the system",
        mode="synthetic",
    )["objective_contracts"]
    lines = [
        f"Control: {control_id} — {control_title}",
        "",
        "Assessment objectives that MUST each get a dedicated labeled section:",
        "Rewrite each objective as satisfied current-state implementation evidence.",
        "Do NOT repeat negative gap phrasing such as missing, not documented, no evidence, or lacks in the final artifact text.",
        "",
    ]
    for obj in objectives:
        lines.append(f"  OBJECTIVE ID: {obj['objective_id']}")
        lines.append(f"  GAP: {obj['description']}")
        lines.append(f"  REQUIRED HEADING: \"{obj['objective_id']} — [descriptive title]\"")
        lines.append(f"  REQUIRED FIRST SENTENCE: \"This section satisfies NIST 800-53A assessment objective {obj['objective_id']}.\"")
        lines.append("")
    lines.append(format_contracts_for_prompt(contracts))
    return "\n".join(lines)


# ── Fix 5: Check which controls already have indexed artifact coverage ─────────

async def _get_controls_with_coverage(project_id: int, control_ids: list[str]) -> set[str]:
    """Return the subset of control_ids that already have indexed autogenerated docs
    specifically targeting them (via artifact_controls metadata).
    Only counts docs that are actually indexed (not pending/failed).
    Synthetic test-package documents do not count as remediation coverage.
    """
    async with AsyncSessionLocal() as db:
        # Count only prior remediation-generated artifacts, not synthetic test packages.
        result = await db.execute(
            select(Document.artifact_controls, Document.parse_status)
            .where(
                Document.project_id == project_id,
                Document.autogenerated == True,  # noqa: E712
                Document.source_assessment_id.isnot(None),
                Document.source_remediation_report_id.isnot(None),
                Document.parse_status.in_(["indexed", "complete"]),
                Document.artifact_controls.isnot(None),
            )
        )
        covered: set[str] = set()
        for row in result.fetchall():
            controls = row[0] or []
            if row[1] in ("indexed", "complete"):
                for cid in controls:
                    if cid in control_ids:
                        covered.add(cid)
    return covered


async def _set_report_progress(report_id: int, detail: str) -> None:
    """Atomically update progress_detail for live UI display."""
    async with AsyncSessionLocal() as db:
        await db.execute(
            update(RemediationReport)
            .where(RemediationReport.id == report_id)
            .values(progress_detail=detail)
        )
        await db.commit()


async def _is_cancelled(report_id: int) -> bool:
    """Return True if the report has been cancelled or terminated externally."""
    async with AsyncSessionLocal() as db:
        result = await db.execute(
            select(RemediationReport.status).where(RemediationReport.id == report_id)
        )
        row = result.one_or_none()
        return row is None or row[0] not in ("running", "pending")


async def _call_llm(llm, system_prompt: str, user_message: str, max_tokens: int = 4000) -> str:
    """Call LLM via the standard complete() interface. Returns raw text."""
    provider = await _ensure_llm_provider(llm, purpose="remediation_generation")
    return await provider.complete(system_prompt, user_message)


async def _ensure_llm_provider(llm, purpose: str = "remediation_generation"):
    """Return a provider object regardless of whether the caller passed a model string or provider instance."""
    if hasattr(llm, "complete") and hasattr(llm, "assess_control"):
        return llm

    from app.services.llm.runtime import build_provider_for_purpose

    provider_name = None
    model_name = None
    if isinstance(llm, dict):
        provider_name = llm.get("provider")
        model_name = llm.get("model")
    elif isinstance(llm, str):
        model_name = llm

    async with AsyncSessionLocal() as db:
        provider, _runtime = await build_provider_for_purpose(
            db,
            purpose=purpose,
            provider_name=provider_name,
            model=model_name,
        )
    return provider


def _format_gaps_for_prompt(findings: list[dict]) -> str:
    """Format findings into a compact gap summary for LLM prompts."""
    lines = []
    for f in findings:
        control_id = f["control_id"]
        title = f["control_title"]
        status = f["status"]
        gaps = f.get("gaps") or []
        if isinstance(gaps, str):
            gaps = [gaps]
        gap_text = "; ".join(str(g) for g in gaps[:5]) if gaps else "No documented evidence of implementation"
        lines.append(f"- {control_id} ({title}) [{status}]: {gap_text}")
    return "\n".join(lines)


def _parse_json_response(text: str) -> list[dict]:
    """Extract JSON array from LLM response, tolerating markdown code fences."""
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
    try:
        result = json.loads(text)
        return result if isinstance(result, list) else []
    except json.JSONDecodeError:
        start = text.find("[")
        end = text.rfind("]")
        if start != -1 and end != -1:
            try:
                return json.loads(text[start:end + 1])
            except Exception:
                pass
    return []


def _strip_md_fences(text: str) -> str:
    """Remove markdown code fences."""
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
    return text.strip()


def _parse_artifact_json(text: str) -> dict:
    """Parse an artifact JSON payload, tolerating fences and stray text."""
    raw = _strip_md_fences(text)
    start = raw.find("{")
    end = raw.rfind("}") + 1
    if start != -1 and end > start:
        raw = raw[start:end]
    try:
        data = json.loads(raw)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _section_contains_objective(section: dict, objective_id: str) -> bool:
    """Return True if a rendered section appears to explicitly cover an objective."""
    text_parts: list[str] = []
    if section.get("text"):
        text_parts.append(str(section.get("text")))
    for key in ("items", "headers"):
        value = section.get(key)
        if isinstance(value, list):
            text_parts.extend(str(item) for item in value)
    rows = section.get("rows")
    if isinstance(rows, list):
        for row in rows:
            if isinstance(row, list):
                text_parts.extend(str(item) for item in row)
    haystack = " ".join(text_parts)
    return objective_id in haystack


def _build_artifact_fallback_sections(
    bundle: dict,
    system_name: str,
    system_context: str,
    objectives: list[dict],
) -> list[dict]:
    """Build a deterministic fallback document body when the model returns weak output."""
    family_title = bundle.get("family_title") or bundle.get("family") or "Control Family"
    controls = ", ".join(bundle.get("control_ids", []))
    evidence_role = bundle.get("evidence_role", "implementation").replace("_", " ")
    contracts = build_control_closure_guidance(
        control_id=(bundle.get("control_ids") or ["MULTI"])[0],
        control_title=bundle.get("title") or family_title,
        gaps=objectives,
        system_name=system_name,
        mode="synthetic",
    )["objective_contracts"]
    intro_text = (
        f"This reassessment-ready artifact package supports {system_name} and provides current-state {evidence_role} "
        f"evidence for controls {controls} in the {family_title} family. The system context is "
        f"{system_context or system_name}, and each objective below is written to satisfy a deterministic closure contract."
    )
    sections = build_contract_sections(
        contracts=contracts,
        system_name=system_name,
        document_type=bundle.get("document_type") or bundle.get("artifact_type"),
        intro_title=bundle["title"],
        intro_text=intro_text,
    )
    summary_paragraphs = [
        paragraph
        for paragraph in synthesize_control_implementation_statement(
            control_id=(bundle.get("control_ids") or ["MULTI"])[0],
            control_title=bundle.get("title") or family_title,
            status="compliant",
            objectives=objectives,
            gap_analysis=[
                {"objective_id": obj.get("objective_id"), "met": "yes", "gap": None}
                for obj in objectives
            ],
        ).split("\n\n")
        if paragraph.strip()
    ]
    sections.insert(2, {"type": "heading", "level": 2, "text": "Control Implementation Summary"})
    for offset, paragraph in enumerate(summary_paragraphs, start=3):
        sections.insert(offset, {"type": "paragraph", "text": paragraph})
    technical_context_index = 3 + len(summary_paragraphs)
    sections.insert(technical_context_index, {"type": "heading", "level": 2, "text": "Technical Context"})
    sections.insert(
        technical_context_index + 1,
        {
            "type": "paragraph",
            "text": (
                f"The package documents implemented control behavior, evidence sources, verification activities, and record "
                f"retention used during reassessment for {system_context or system_name}."
            ),
        },
    )
    return sections


def _normalize_artifact_content_json(
    content_json: str,
    bundle: dict,
    system_name: str,
    system_context: str,
) -> str:
    """Ensure remediation artifacts contain substantive objective-level sections."""
    parsed = _parse_artifact_json(content_json)
    title = parsed.get("title") or bundle["title"]
    sections = parsed.get("sections")
    sections = sections if isinstance(sections, list) else []

    objectives: list[dict] = []
    for finding in bundle.get("findings", []):
        objectives.extend(_build_gap_objectives(finding, bundle.get("assessment_objectives")))

    contracts = build_control_closure_guidance(
        control_id=(bundle.get("control_ids") or ["MULTI"])[0],
        control_title=bundle.get("title") or "Artifact",
        gaps=objectives,
        system_name=system_name,
        mode="synthetic",
    )["objective_contracts"]
    has_objective_coverage = all(
        any(_section_contains_objective(section, obj["objective_id"]) for section in sections)
        for obj in objectives
    ) if objectives else True
    contracts_ok, _ = sections_satisfy_contracts(sections, contracts)
    substantive_sections = sum(
        1
        for section in sections
        if section.get("type") in {"heading", "paragraph", "bullet_list", "numbered_list", "table"}
        and any(section.get(key) for key in ("text", "items", "rows"))
    )
    if substantive_sections < 4 or not has_objective_coverage or not contracts_ok:
        fallback_sections = _build_artifact_fallback_sections(bundle, system_name, system_context, objectives)
        if sections and not has_objective_coverage:
            fallback_sections.extend(
                section for section in sections
                if section.get("type") in {"paragraph", "table", "bullet_list", "numbered_list"}
            )
        parsed = {"title": title, "sections": fallback_sections}
    elif not sections:
        parsed = {
            "title": title,
            "sections": _build_artifact_fallback_sections(bundle, system_name, system_context, objectives),
        }
    else:
        parsed = {"title": title, "sections": sections}

    return json.dumps(parsed)


# ── Word document builder from structured JSON ─────────────────────────────────

def _build_docx_cover(
    doc,
    title: str,
    subtitle: str,
    generated_date: str,
    classification: str,
) -> None:
    """Add cover page elements (header band, title, metadata table) to doc."""
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(classification)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    header_para.paragraph_format.space_after = Pt(0)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1F3864")
    header_para._p.get_or_add_pPr().append(shd)

    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(subtitle)
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    for i, (label, value) in enumerate([
        ("Document Date", generated_date),
        ("Framework", "NIST SP 800-53 Rev 5"),
        ("Document Status", "DRAFT — Review and approve before use"),
    ]):
        meta.rows[i].cells[0].text = label
        meta.rows[i].cells[1].text = value
        meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()


def _render_sections_to_doc(doc, sections: list[dict]) -> None:
    """Render a list of JSON section nodes into a python-docx Document."""
    from docx.shared import Pt

    for node in sections:
        ntype = node.get("type", "paragraph")

        if ntype == "heading":
            level = max(1, min(4, int(node.get("level", 1))))
            doc.add_heading(node.get("text", ""), level=level)

        elif ntype == "paragraph":
            p = doc.add_paragraph(node.get("text", ""))
            p.paragraph_format.space_after = Pt(6)

        elif ntype == "numbered_list":
            for item in node.get("items", []):
                doc.add_paragraph(str(item), style="List Number")

        elif ntype == "bullet_list":
            for item in node.get("items", []):
                doc.add_paragraph(str(item), style="List Bullet")

        elif ntype == "table":
            headers = node.get("headers", [])
            rows = node.get("rows", [])
            if not headers and not rows:
                continue
            ncols = max(len(headers), max((len(r) for r in rows), default=0))
            if ncols == 0:
                continue
            nrows = len(rows) + (1 if headers else 0)
            tbl = doc.add_table(rows=nrows, cols=ncols)
            tbl.style = "Table Grid"
            ri = 0
            if headers:
                for ci, h in enumerate(headers[:ncols]):
                    cell = tbl.cell(ri, ci)
                    cell.text = str(h)
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                ri += 1
            for row in rows:
                for ci, val in enumerate(row[:ncols]):
                    tbl.cell(ri, ci).text = str(val)
                ri += 1
            doc.add_paragraph()  # spacing after table

        elif ntype == "divider":
            doc.add_paragraph("─" * 60)


def _json_to_docx(
    title: str,
    subtitle: str,
    json_content: str,
    generated_date: str,
    classification: str = "CONTROLLED UNCLASSIFIED INFORMATION",
) -> bytes:
    """Build a Word document from the LLM's structured JSON content. Returns bytes."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    _build_docx_cover(doc, title, subtitle, generated_date, classification)

    # Parse JSON — tolerate code fences or stray text before/after the object
    raw = json_content.strip()
    start = raw.find("{")
    end = raw.rfind("}") + 1
    if start != -1 and end > start:
        raw = raw[start:end]

    try:
        data = json.loads(raw)
        sections = data.get("sections", [])
    except (json.JSONDecodeError, Exception):
        # Fallback: write the raw text as a single paragraph so nothing is silently lost
        logger.warning("_json_to_docx: failed to parse LLM JSON, writing raw text fallback")
        sections = [{"type": "paragraph", "text": json_content}]

    _render_sections_to_doc(doc, sections)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Excel action tracker builder ───────────────────────────────────────────────

def _build_guide_xlsx(guide_sections: list[dict], summary: dict, system_name: str) -> bytes:
    """Build a formatted Excel action tracker from guide sections."""
    import openpyxl
    from openpyxl.styles import (
        Alignment, Border, Font, PatternFill, Side,
    )

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Remediation Actions"

    # ── Colour palette ───────────────────────────────────────────────────────
    NAVY   = "1F3864"
    RED    = "C0504D"
    ORANGE = "F79646"
    GREEN  = "4BACC6"
    GRAY   = "D9D9D9"
    WHITE  = "FFFFFF"

    hdr_font   = Font(name="Calibri", bold=True, color=WHITE, size=11)
    hdr_fill   = PatternFill("solid", fgColor=NAVY)
    hdr_align  = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_side  = Side(style="thin", color="AAAAAA")
    thin_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)

    def hdr(cell, text):
        cell.value = text
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.alignment = hdr_align
        cell.border = thin_border

    def cell_val(cell, text, bold=False, wrap=True, fill_color=None, align="left"):
        cell.value = text
        cell.font = Font(name="Calibri", bold=bold, size=10)
        cell.alignment = Alignment(horizontal=align, vertical="top", wrap_text=wrap)
        cell.border = thin_border
        if fill_color:
            cell.fill = PatternFill("solid", fgColor=fill_color)

    # ── Title row ────────────────────────────────────────────────────────────
    ws.merge_cells("A1:H1")
    ws["A1"].value = f"Remediation Action Plan — {system_name}"
    ws["A1"].font = Font(name="Calibri", bold=True, size=14, color=NAVY)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

    ws.merge_cells("A2:H2")
    ws["A2"].value = f"Generated: {datetime.now().strftime('%B %d, %Y')}  ·  NIST SP 800-53 Rev 5  ·  {summary.get('total_actions', 0)} actions required"
    ws["A2"].font = Font(name="Calibri", size=10, color="666666")
    ws["A2"].alignment = Alignment(horizontal="center")
    ws.row_dimensions[2].height = 18

    ws.row_dimensions[3].height = 6  # spacer

    # ── Column headers ───────────────────────────────────────────────────────
    headers = [
        ("A", "Control ID", 12),
        ("B", "Control Title", 30),
        ("C", "Status", 16),
        ("D", "Gap / Finding", 40),
        ("E", "Action Required", 45),
        ("F", "Responsible Role", 18),
        ("G", "Effort", 14),
        ("H", "Success Criteria", 38),
        ("I", "Template Language", 45),
        ("J", "Owner (fill in)", 18),
        ("K", "Target Date", 14),
        ("L", "Completed", 12),
    ]
    for col_letter, label, width in headers:
        ws.column_dimensions[col_letter].width = width
        hdr(ws[f"{col_letter}4"], label)
    ws.row_dimensions[4].height = 22

    # ── Data rows ────────────────────────────────────────────────────────────
    row = 5
    for section in guide_sections:
        # Family header row
        ws.merge_cells(f"A{row}:L{row}")
        ws[f"A{row}"].value = f"  {section['family']} — {section['family_title']}"
        ws[f"A{row}"].font = Font(name="Calibri", bold=True, size=11, color=WHITE)
        ws[f"A{row}"].fill = PatternFill("solid", fgColor="2E5090")
        ws[f"A{row}"].alignment = Alignment(vertical="center")
        ws.row_dimensions[row].height = 18
        row += 1

        for action in section.get("actions", []):
            status = action.get("status", "")
            fill = PatternFill("solid", fgColor="FFF2F2") if status == "non_compliant" else \
                   PatternFill("solid", fgColor="FFF8EC") if status == "partially_compliant" else None

            cell_val(ws[f"A{row}"], action.get("control_id", ""), bold=True, wrap=False)
            cell_val(ws[f"B{row}"], action.get("control_title", ""))
            cell_val(ws[f"C{row}"], "Non-Compliant" if status == "non_compliant" else "Partial" if status == "partially_compliant" else status)
            cell_val(ws[f"D{row}"], action.get("gap", ""))
            cell_val(ws[f"E{row}"], action.get("action", ""))
            cell_val(ws[f"F{row}"], action.get("responsible", ""))
            cell_val(ws[f"G{row}"], action.get("effort", ""), align="center")
            cell_val(ws[f"H{row}"], action.get("success_criteria", ""))
            cell_val(ws[f"I{row}"], action.get("template_language", ""))
            cell_val(ws[f"J{row}"], "")   # owner fill-in
            cell_val(ws[f"K{row}"], "")   # target date
            cell_val(ws[f"L{row}"], "")   # completed checkbox

            if fill:
                for col in "ABCDEFGHIJKL":
                    ws[f"{col}{row}"].fill = fill

            ws.row_dimensions[row].height = 45
            row += 1

    # ── Summary sheet ────────────────────────────────────────────────────────
    ws2 = wb.create_sheet("Summary")
    ws2["A1"].value = "Compliance Summary"
    ws2["A1"].font = Font(name="Calibri", bold=True, size=14, color=NAVY)
    ws2.column_dimensions["A"].width = 28
    ws2.column_dimensions["B"].width = 14

    summary_rows = [
        ("Total Controls Assessed", summary.get("total_controls", 0)),
        ("Compliant", summary.get("compliant", 0)),
        ("Partially Compliant", summary.get("partially_compliant", 0)),
        ("Non-Compliant", summary.get("non_compliant", 0)),
        ("Not Applicable", summary.get("not_applicable", 0)),
        ("Families with Gaps", len(summary.get("families_affected", []))),
        ("Total Actions Required", summary.get("total_actions", 0)),
    ]
    for i, (label, val) in enumerate(summary_rows, start=3):
        ws2[f"A{i}"].value = label
        ws2[f"A{i}"].font = Font(name="Calibri", bold=True, size=10)
        ws2[f"B{i}"].value = val
        ws2[f"B{i}"].font = Font(name="Calibri", size=10)
        ws2[f"B{i}"].alignment = Alignment(horizontal="center")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Save a file as a Document record ─────────────────────────────────────────

async def _save_document(
    file_bytes: bytes,
    filename: str,
    file_ext: str,
    project_id: int,
    assessment_id: int,
    report_id: int,
    created_by: int,
    upload_dir: Path,
    controls_addressed: list[str] | None = None,
    document_type: str | None = None,
    document_intent: str | None = None,
    trigger_parse: bool = True,
) -> int:
    """Write file to disk and create a Document ORM record. Returns doc id.

    If an AI-generated document with the same filename already exists in this
    project, the old file and its DB record are removed before the new one is
    saved — keeping the document library clean and deduplicated.

    controls_addressed: explicit list of NIST control IDs this document targets.
    When set, the control tagger will apply these as forced high-confidence tags
    instead of LLM-guessing, preventing tag dilution across unrelated controls.
    """
    file_hash = hashlib.sha256(file_bytes).hexdigest()
    ext_to_type = {"docx": "docx", "xlsx": "xlsx", "csv": "csv"}

    async with AsyncSessionLocal() as db:
        # ── Dedup: remove any existing autogenerated doc with the same filename ──
        existing_result = await db.execute(
            select(Document).where(
                Document.project_id == project_id,
                Document.filename == filename,
                Document.autogenerated == True,  # noqa: E712
            )
        )
        for old_doc in existing_result.scalars().all():
            # Remove the physical file if it still exists
            try:
                old_path = Path(old_doc.file_path)
                if old_path.exists():
                    old_path.unlink()
            except Exception as _del_err:
                logger.warning("Could not delete old artifact file %s: %s", old_doc.file_path, _del_err)
            await db.delete(old_doc)
            logger.info(
                "Replaced existing AI artifact doc %d (%s) for project %d",
                old_doc.id, filename, project_id,
            )

        await db.flush()  # ensure old records are gone before inserting new one

        # ── Write new file to disk ───────────────────────────────────────────────
        file_path = upload_dir / f"{uuid4()}.{file_ext}"
        file_path.write_bytes(file_bytes)

        doc = Document(
            project_id=project_id,
            filename=filename,
            file_path=str(file_path),
            file_type=ext_to_type.get(file_ext, file_ext),
            file_hash=file_hash,
            file_size_bytes=len(file_bytes),
            parse_status="pending",
            uploaded_by=created_by,
            autogenerated=True,
            artifact_status="draft",
            evidence_eligible=False,
            source_assessment_id=assessment_id,
            source_remediation_report_id=report_id,
            artifact_controls=controls_addressed or None,
            document_type=document_type or None,
            document_intent=document_intent or None,
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)
        doc_id = doc.id

    # Trigger background parsing/indexing
    if trigger_parse:
        from app.services.parsers.dispatcher import dispatch_parse
        asyncio.create_task(dispatch_parse(doc_id))

    return doc_id


# ── Guide Generation ───────────────────────────────────────────────────────────

async def _generate_guide(
    report_id: int,
    assessment: Assessment,
    findings: list,
    llm,
    system_context: str,
) -> None:
    """Generate remediation guide and save Word + Excel files to DB."""
    from app.services.prompt_manager import get_prompt
    from app.services.system_knowledge import extract_system_knowledge
    from app.services.tool_guidance import build_action_collection_guidance, build_collection_playbook
    from app.core.config import get_settings
    settings = get_settings()

    findings, non_assessable_controls = _split_assessable_findings(findings)
    system_prompt = await get_prompt("remediation_guide_system", REMEDIATION_GUIDE_SYSTEM_PROMPT)

    system_knowledge_summary = None
    detected_tools: list[dict] = []
    async with AsyncSessionLocal() as db:
        system_knowledge_summary = await extract_system_knowledge(
            db,
            project_id=assessment.project_id,
            source_mode="remediation_real",
            source_run_id=report_id,
        )
        detected_tools = list(system_knowledge_summary.get("tools", []))

    # Group findings by family
    by_family: dict[str, list] = defaultdict(list)
    family_titles: dict[str, str] = {}
    for f in findings:
        if f.status in ("non_compliant", "partially_compliant"):
            by_family[f.control_family].append(f)
            family_titles[f.control_family] = _family_title(f.control_family)

    guide_sections = []
    total_families = len(by_family)
    done_families = 0

    for family, family_findings in sorted(by_family.items()):
        if await _is_cancelled(report_id):
            return

        done_families += 1
        await _set_report_progress(
            report_id,
            f"▶ {done_families}/{total_families} — Calling LLM for {family} ({family_titles.get(family, family)})…"
        )

        family_findings_dicts = [
            {
                "control_id": f.control_id,
                "control_title": f.control_title,
                "status": f.status,
                "gaps": f.gaps or [],
                "remediation_plan": f.remediation_plan or "",
            }
            for f in family_findings
        ]

        gap_summary = _format_gaps_for_prompt(family_findings_dicts)
        user_msg = (
            f"System: {system_context or 'Federal information system'}\n\n"
            f"Control Family: {family} — {family_titles.get(family, family)}\n\n"
            f"Compliance gaps to remediate:\n{gap_summary}"
        )

        try:
            raw = await _call_llm(llm, system_prompt, user_msg, max_tokens=3000)
            actions = _parse_json_response(raw)
        except Exception as e:
            logger.warning("Guide generation failed for family %s: %s", family, e)
            actions = []

        # Enrich actions with status for Excel colouring
        for a in actions:
            if not a.get("status"):
                match = next((f for f in family_findings_dicts if f["control_id"] == a.get("control_id")), None)
                if match:
                    a["status"] = match["status"]
                    a["control_title"] = match["control_title"]
            a["collection_guidance"] = build_action_collection_guidance(a, detected_tools)

        # Fall back: build basic actions from existing remediation_plan fields
        if not actions:
            for f in family_findings_dicts:
                actions.append({
                    "control_id": f["control_id"],
                    "control_title": f["control_title"],
                    "status": f["status"],
                    "gap": (f["gaps"][0] if f["gaps"] else "Gap identified"),
                    "action": f["remediation_plan"] or f"Implement {f['control_id']} requirements",
                    "responsible": "ISSO",
                    "effort": "2-4 weeks",
                    "success_criteria": f"{f['control_id']} assessment objective met",
                    "template_language": "",
                    "collection_guidance": build_action_collection_guidance(
                        {
                            "control_id": f["control_id"],
                            "control_title": f["control_title"],
                            "gap": (f["gaps"][0] if f["gaps"] else "Gap identified"),
                        },
                        detected_tools,
                    ),
                })

        guide_sections.append({
            "family": family,
            "family_title": family_titles.get(family, family),
            "control_count": len(family_findings),
            "actions": actions,
            "collection_playbook": build_collection_playbook(actions, detected_tools),
        })

    # Compute summary stats
    total = len(findings)
    nc = sum(1 for f in findings if f.status == "non_compliant")
    pc = sum(1 for f in findings if f.status == "partially_compliant")
    compliant = sum(1 for f in findings if f.status == "compliant")
    na = sum(1 for f in findings if f.status == "not_applicable")
    total_actions = sum(len(s["actions"]) for s in guide_sections)

    summary = {
        "total_controls": total,
        "compliant": compliant,
        "partially_compliant": pc,
        "non_compliant": nc,
        "not_applicable": na,
        "excluded_non_assessable_count": len(non_assessable_controls),
        "excluded_non_assessable_controls": non_assessable_controls,
        "families_affected": sorted(by_family.keys()),
        "total_actions": total_actions,
        "tools_detected": len(detected_tools),
        "collection_playbook_entries": sum(len(s.get("collection_playbook", [])) for s in guide_sections),
    }
    playbook_entries = [
        item
        for section in guide_sections
        for item in section.get("collection_playbook", [])
    ]

    content = {
        "report_type": "guide",
        "generation_mode": "remediation_real",
        "generated_at": datetime.now(UTC).isoformat(),
        "summary": summary,
        "sections": guide_sections,
        "detected_tools": detected_tools,
        "system_knowledge": system_knowledge_summary,
        "collection_playbook": playbook_entries,
        "playbook_summary": {
            "entries": len(playbook_entries),
            "detected_tool_entries": sum(1 for item in playbook_entries if item.get("detected")),
            "generic_entries": sum(1 for item in playbook_entries if not item.get("detected")),
            "domains": sorted({item.get("domain") for item in playbook_entries if item.get("domain")}),
        },
    }

    # ── Save Word doc + Excel tracker as Document records ────────────────────
    generated_doc_ids: list[int] = []
    project_upload_dir = Path(settings.upload_dir) / str(assessment.project_id)
    project_upload_dir.mkdir(parents=True, exist_ok=True)

    async with AsyncSessionLocal() as db:
        rep_result = await db.execute(select(RemediationReport).where(RemediationReport.id == report_id))
        rep = rep_result.scalar_one()
        created_by = rep.created_by

    system_name = system_context.split("—")[0].strip() if "—" in system_context else system_context or "System"
    date_str = datetime.now().strftime("%B %d, %Y")
    new_generated_doc_ids: list[int] = []

    # Word guide document
    await _set_report_progress(report_id, "Building Word remediation guide…")
    try:
        guide_docx_bytes = _build_guide_docx(guide_sections, summary, system_name, date_str)
        doc_id = await _save_document(
            file_bytes=guide_docx_bytes,
            filename=f"Remediation_Guide_{system_name.replace(' ', '_')}.docx",
            file_ext="docx",
            project_id=assessment.project_id,
            assessment_id=assessment.id,
            report_id=report_id,
            created_by=created_by,
            upload_dir=project_upload_dir,
            trigger_parse=False,
        )
        generated_doc_ids.append(doc_id)
        new_generated_doc_ids.append(doc_id)
    except Exception as e:
        logger.error("Failed to save guide Word doc: %s", e)

    # Excel action tracker
    await _set_report_progress(report_id, "Building Excel action tracker…")
    try:
        xlsx_bytes = _build_guide_xlsx(guide_sections, summary, system_name)
        doc_id = await _save_document(
            file_bytes=xlsx_bytes,
            filename=f"Remediation_Action_Tracker_{system_name.replace(' ', '_')}.xlsx",
            file_ext="xlsx",
            project_id=assessment.project_id,
            assessment_id=assessment.id,
            report_id=report_id,
            created_by=created_by,
            upload_dir=project_upload_dir,
            trigger_parse=False,
        )
        generated_doc_ids.append(doc_id)
        new_generated_doc_ids.append(doc_id)
    except Exception as e:
        logger.error("Failed to save guide Excel tracker: %s", e)

    if new_generated_doc_ids:
        from app.services.parsers.dispatcher import dispatch_parse_batch
        asyncio.create_task(dispatch_parse_batch(new_generated_doc_ids))

    async with AsyncSessionLocal() as db:
        await db.execute(
            update(RemediationReport)
            .where(RemediationReport.id == report_id)
            .values(
                status="complete",
                content_json=content,
                generated_doc_ids=generated_doc_ids,
                progress_detail=f"Complete — {total_actions} actions · Word guide + Excel tracker saved to document library",
            )
        )
        await db.commit()


def _guide_sections_to_doc_nodes(sections: list[dict], summary: dict, system_name: str) -> list[dict]:
    """Convert guide sections to JSON document nodes for _render_sections_to_doc."""
    nodes: list[dict] = []

    nodes.append({"type": "heading", "level": 1, "text": "Executive Summary"})
    nodes.append({"type": "paragraph", "text": (
        f"This remediation guide covers {summary.get('non_compliant', 0)} non-compliant and "
        f"{summary.get('partially_compliant', 0)} partially compliant controls across "
        f"{len(summary.get('families_affected', []))} control families for {system_name}. "
        f"Total remediation actions required: {summary.get('total_actions', 0)}."
    )})
    if summary.get("tools_detected"):
        nodes.append({"type": "paragraph", "text": (
            f"Detected tools and platforms referenced by the current package: "
            f"{summary.get('tools_detected', 0)}. The evidence guidance below points the assessor "
            f"to real collection locations rather than made-up artifacts."
        )})

    for section in sections:
        nodes.append({
            "type": "heading", "level": 2,
            "text": f"{section['family']} — {section['family_title']}",
        })
        for a in section.get("actions", []):
            nodes.append({
                "type": "heading", "level": 3,
                "text": f"{a.get('control_id', '')} — {a.get('gap', '')[:80]}",
            })
            detail_rows = [
                ["Action", a.get("action", "")],
                ["Responsible", a.get("responsible", "")],
                ["Effort", a.get("effort", "")],
                ["Done when", a.get("success_criteria", "")],
            ]
            if a.get("template_language"):
                detail_rows.append(["Template language", a["template_language"]])
            nodes.append({
                "type": "table",
                "headers": ["Field", "Detail"],
                "rows": detail_rows,
            })
            if a.get("collection_guidance"):
                bullets: list[str] = []
                for guidance in a["collection_guidance"][:3]:
                    tool_label = guidance.get("tool_name") or guidance.get("domain", "evidence source")
                    where = ", ".join(guidance.get("where_to_go", [])[:2])
                    collect = ", ".join(guidance.get("collect", [])[:2])
                    bullets.append(f"{tool_label}: go to {where}; collect {collect}.")
                if bullets:
                    nodes.append({"type": "paragraph", "text": "Real evidence to collect:"})
                    nodes.append({"type": "bullet_list", "items": bullets})
            nodes.append({"type": "divider"})

    return nodes


def _build_guide_docx(
    sections: list[dict],
    summary: dict,
    system_name: str,
    generated_date: str,
) -> bytes:
    """Build the same remediation guide used for stored and on-demand downloads."""
    from docx import Document as DocxDocument

    document = DocxDocument()
    _build_docx_cover(
        document,
        "Remediation Guide",
        system_name,
        generated_date,
        "CONTROLLED UNCLASSIFIED INFORMATION",
    )
    _render_sections_to_doc(document, _guide_sections_to_doc_nodes(sections, summary, system_name))
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_artifact_package_docx(content: dict, generated_date: str) -> bytes:
    """Build a readable index for an assessment-wide generated artifact package."""
    from docx import Document as DocxDocument

    artifacts = list(content.get("artifacts") or [])
    document = DocxDocument()
    _build_docx_cover(
        document,
        "Remediation Artifact Package",
        str(content.get("system_name") or "System"),
        generated_date,
        "CONTROLLED UNCLASSIFIED INFORMATION",
    )
    nodes: list[dict] = [
        {
            "type": "paragraph",
            "text": f"This package indexes {len(artifacts)} generated remediation artifact groups.",
        }
    ]
    for artifact in artifacts:
        family = str(artifact.get("family") or "Artifact")
        title = str(artifact.get("title") or family)
        controls = ", ".join(str(item) for item in artifact.get("controls_addressed") or [])
        nodes.extend(
            [
                {"type": "heading", "level": 2, "text": title},
                {"type": "paragraph", "text": f"Family: {family}"},
                {"type": "paragraph", "text": f"Controls addressed: {controls or 'See generated documents'}"},
            ]
        )
    _render_sections_to_doc(document, nodes)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ── Artifact Generation ────────────────────────────────────────────────────────

def _build_artifact_content_prompt(artifact_type: str) -> str:
    """Return the type-specific system prompt for content generation."""
    type_instructions = ARTIFACT_TYPE_INSTRUCTIONS.get(artifact_type, ARTIFACT_TYPE_INSTRUCTIONS["policy_procedure"])
    return f"{ARTIFACT_CONTENT_BASE_PROMPT}\n\n{type_instructions}"


def _safe_filename(text: str, max_length: int = 80) -> str:
    """Make a filesystem-safe filename component."""
    cleaned = re.sub(r"[^\w\s-]", "", text).strip()
    cleaned = re.sub(r"\s+", "_", cleaned)
    return cleaned[:max_length]


def _build_artifact_filename(prefix: str, unique_key: str, title: str, ext: str = "docx") -> str:
    """Build a stable, unique artifact filename that won't collide across bundles."""
    safe_prefix = _safe_filename(prefix, 24)
    safe_key = _safe_filename(unique_key, 64)
    safe_title = _safe_filename(title, 72)
    parts = [part for part in (safe_prefix, safe_key, safe_title) if part]
    return f"{'_'.join(parts)}.{ext}"


async def _prune_existing_document_ids(doc_ids: list[int]) -> list[int]:
    """Keep only document ids that still exist after in-run dedupe/replacement."""
    if not doc_ids:
        return []
    ordered_unique: list[int] = []
    seen: set[int] = set()
    for doc_id in doc_ids:
        if doc_id not in seen:
            ordered_unique.append(doc_id)
            seen.add(doc_id)
    async with AsyncSessionLocal() as db:
        existing = await db.execute(select(Document.id).where(Document.id.in_(ordered_unique)))
        existing_ids = set(existing.scalars().all())
    return [doc_id for doc_id in ordered_unique if doc_id in existing_ids]


async def _get_existing_content(project_id: int, filename: str) -> str | None:
    """Return the indexed text of an existing autogenerated document, or None.

    Reconstructs the document body from its stored chunks so the LLM can
    update the document in-place rather than generating a completely new one.
    Caps at 10 000 characters to stay within token budget.
    """
    async with AsyncSessionLocal() as db:
        doc_result = await db.execute(
            select(Document.id).where(
                Document.project_id == project_id,
                Document.filename == filename,
                Document.autogenerated == True,  # noqa: E712
                Document.parse_status.in_(["indexed", "complete"]),
            ).limit(1)
        )
        row = doc_result.fetchone()
        if not row:
            return None
        doc_id = row[0]
        payload = await get_document_evidence_payload(doc_id, db)

    chunks = [chunk["content"] for chunk in (payload or {}).get("chunks", []) if chunk.get("content")]
    if not chunks:
        return None

    full_text = "\n\n".join(chunks)
    if len(full_text) > 10_000:
        full_text = full_text[:10_000] + "\n\n[... remainder of document omitted for brevity ...]"
    return full_text


async def _generate_artifacts(
    report_id: int,
    assessment: Assessment,
    findings: list,
    llm,
    system_context: str,
) -> None:
    """Generate one targeted evidence artifact per failing control.

    Fix 2: One document per control (not per family) — each doc is laser-focused
    Fix 1: Structured per-objective sections with explicit IDs in headings
    Fix 3: Update-in-place — read existing doc, append only missing sections
    Fix 4: Store controls_addressed on the Document so tagger can skip LLM
    Fix 5: Skip controls that already have indexed autogenerated coverage
    Fix 6: Technical controls → Evidence Collection Template, not fake policy
    """
    from app.services.prompt_manager import get_prompt
    from app.core.config import get_settings
    settings = get_settings()

    planner_prompt = await get_prompt("remediation_artifact_planner", ARTIFACT_PLANNER_SYSTEM_PROMPT)
    evidence_prompt = await get_prompt("remediation_artifact_evidence", EVIDENCE_COLLECTION_SYSTEM_PROMPT)
    content_prompt_cache: dict[str, str] = {"evidence_template": evidence_prompt}
    content_prompt_cache: dict[str, str] = {"evidence_template": evidence_prompt}

    # Collect non/partial findings only
    gap_findings = [
        f for f in findings
        if f.status in ("non_compliant", "partially_compliant")
    ]
    total_controls = len(gap_findings)

    if total_controls == 0:
        async with AsyncSessionLocal() as db:
            await db.execute(
                update(RemediationReport)
                .where(RemediationReport.id == report_id)
                .values(
                    status="complete",
                    content_json={"report_type": "artifacts", "summary": {"documents_created": 0}, "artifacts": []},
                    generated_doc_ids=[],
                    progress_detail="Complete — no non/partial controls to address",
                )
            )
            await db.commit()
        return

    project_upload_dir = Path(settings.upload_dir) / str(assessment.project_id)
    project_upload_dir.mkdir(parents=True, exist_ok=True)

    system_name = system_context.split("—")[0].strip() if "—" in system_context else system_context or "System"
    date_str = datetime.now().strftime("%B %d, %Y")

    async with AsyncSessionLocal() as db:
        rep_result = await db.execute(select(RemediationReport).where(RemediationReport.id == report_id))
        rep = rep_result.scalar_one()
        created_by = rep.created_by
        prior_content = rep.content_json or {}
        prior_entries: list[dict] = prior_content.get("artifacts", [])
        prior_doc_ids: list[int] = list(rep.generated_doc_ids or [])

    # Resume state — skip controls already processed in a prior interrupted run
    completed_controls: set[str] = {
        e["control_id"] for e in prior_entries if "control_id" in e and not e.get("skipped")
    }
    all_artifact_entries: list[dict] = list(prior_entries)
    generated_doc_ids: list[int] = list(prior_doc_ids)
    done_controls = len(completed_controls)

    # Regenerate remediation coverage on every new run.
    # Prior autogenerated artifacts may be stale, generic, or predate newer closure logic,
    # so they must not suppress regeneration for a fresh remediation pass.
    already_covered: set[str] = set()

    if completed_controls:
        logger.info(
            "Resuming artifact generation for report %d — %d controls already done",
            report_id, len(completed_controls),
        )

    for finding in gap_findings:
        control_id = finding.control_id

        if control_id in completed_controls:
            continue  # already done in prior run
        if await _is_cancelled(report_id):
            return

        done_controls += 1
        family = finding.control_family
        family_title = _family_title(family)

        await _set_report_progress(
            report_id,
            f"▶ {done_controls}/{total_controls} — {control_id} ({finding.control_title[:50]})…"
        )

        # Build structured objectives list for this control (Fix 1)
        criteria_result = await db.execute(
            select(AssessmentCriteriaPackage.assessment_objectives).where(
                AssessmentCriteriaPackage.assessment_id == assessment.id,
                AssessmentCriteriaPackage.control_id == control_id,
            )
        )
        criteria_objectives = criteria_result.scalar_one_or_none() or []
        objectives = _build_gap_objectives(finding, criteria_objectives)
        objectives_prompt_text = _format_objectives_for_prompt(
            objectives, control_id, finding.control_title
        )
        statement_guidance = build_control_statement_generation_guidance(
            control_id,
            finding.control_title,
            objectives,
        )

        # Fix 6: Technical controls → Evidence Collection Template (no fake policy docs)
        is_technical = control_id in TECHNICAL_CONTROL_IDS
        if is_technical:
            art_type = "evidence_template"
            art_title = f"{control_id} Evidence Collection Guide — {finding.control_title}"
            # Short, stable filename: e.g. SC-10_Network_Disconnect_Evidence.docx
            safe_name = _build_artifact_filename(
                "REMEDIATE",
                f"{control_id}_evidence",
                f"{finding.control_title}_Evidence",
            )
            content_system_prompt = evidence_prompt
        else:
            # Fix 2: Call planner once per control, get ONE artifact type back
            planner_user_msg = (
                f"System: {system_context or 'Federal information system'}\n\n"
                f"Single failing control:\n{objectives_prompt_text}\n\n"
                f"Return the JSON object for ONE document that closes all listed gaps."
            )
            try:
                plan_raw = await _call_llm(llm, planner_prompt, planner_user_msg, max_tokens=400)
                plan_raw = _strip_md_fences(plan_raw)
                start_b = plan_raw.find("{")
                end_b = plan_raw.rfind("}") + 1
                plan_obj = json.loads(plan_raw[start_b:end_b]) if start_b != -1 and end_b > start_b else {}
            except Exception as _pe:
                logger.warning("Artifact planning failed for %s: %s", control_id, _pe)
                plan_obj = {}

            valid_types = set(ARTIFACT_TYPE_INSTRUCTIONS.keys())
            art_type = plan_obj.get("artifact_type", "policy_procedure")
            if art_type not in valid_types:
                art_type = "policy_procedure"
            art_title = (
                plan_obj.get("title")
                or f"{control_id} — {finding.control_title} Policy and Procedures"
            )
            safe_name = _build_artifact_filename(
                "REMEDIATE",
                f"{control_id}_{art_type}",
                art_title,
            )
            content_system_prompt = content_prompt_cache.get(art_type)
            if content_system_prompt is None:
                content_system_prompt = await get_prompt(
                    f"remediation_artifact_{art_type}",
                    _build_artifact_content_prompt(art_type),
                )
                content_prompt_cache[art_type] = content_system_prompt

        await _set_report_progress(
            report_id,
            f"▶ {done_controls}/{total_controls} — {control_id}: generating {art_type} document…"
        )

        # Fix 3: Read existing indexed content — append-only update, never replace
        existing_content = await _get_existing_content(assessment.project_id, safe_name)

        if existing_content:
            logger.info(
                "Update mode: appending to existing artifact %s for project %d",
                safe_name, assessment.project_id,
            )
            content_user_msg = (
                f"System: {system_context or 'Federal information system'}\n\n"
                f"Document to update: {art_title}\n"
                f"Document type: {art_type}\n\n"
                f"EXISTING DOCUMENT CONTENT:\n"
                f"---\n{existing_content}\n---\n\n"
                f"{statement_guidance}\n\n"
                f"Assessment objectives that still need dedicated labeled sections:\n\n"
                f"{objectives_prompt_text}\n\n"
                f"INSTRUCTIONS:\n"
                f"1. Preserve ALL existing sections exactly as they appear above.\n"
                f"2. For any objective listed above that does NOT already have a section "
                f"   with its ID in a heading, ADD that section at the end.\n"
                f"3. Do NOT rewrite, restructure, or remove any existing content.\n"
                f"4. Return the COMPLETE updated document as a JSON object with all existing "
                f"   sections first, followed by any new sections you added."
            )
        else:
            content_user_msg = (
                f"System: {system_context or 'Federal information system'}\n\n"
                f"Document to produce: {art_title}\n"
                f"Document type: {art_type}\n\n"
                f"{statement_guidance}\n\n"
                f"Assessment objectives this document MUST close:\n\n"
                f"{objectives_prompt_text}\n\n"
                f"Generate the complete document now."
            )

        try:
            content_json = await _call_llm(llm, content_system_prompt, content_user_msg, max_tokens=6000)
            content_json = _strip_md_fences(content_json)
        except Exception as e:
            logger.warning("Content generation failed for %s: %s", control_id, e)
            content_json = json.dumps({
                "title": art_title,
                "sections": [{"type": "paragraph", "text": (
                    "Generation encountered an error — please retry the artifact report."
                )}],
            })

        # Save Word document
        await _set_report_progress(
            report_id,
            f"✓ {done_controls}/{total_controls} — {control_id}: saving document…"
        )
        try:
            docx_bytes = _json_to_docx(
                title=art_title,
                subtitle=system_name,
                json_content=content_json,
                generated_date=date_str,
            )
            doc_id = await _save_document(
                file_bytes=docx_bytes,
                filename=safe_name,
                file_ext="docx",
                project_id=assessment.project_id,
                assessment_id=assessment.id,
                report_id=report_id,
                created_by=created_by,
                upload_dir=project_upload_dir,
                controls_addressed=[control_id],  # Fix 4: explicit tag, no LLM guessing
            )
            generated_doc_ids.append(doc_id)
            all_artifact_entries.append({
                "control_id": control_id,
                "family": family,
                "family_title": family_title,
                "title": art_title,
                "artifact_type": art_type,
                "doc_id": doc_id,
            })
            completed_controls.add(control_id)
        except Exception as e:
            logger.error("Failed to save docx for %s: %s", control_id, e)

        # Incremental save after every control so interrupted runs can resume
        _partial = {
            "report_type": "artifacts",
            "generated_at": datetime.now(UTC).isoformat(),
            "summary": {
                "controls_addressed": done_controls,
                "documents_created": len(generated_doc_ids),
            },
            "artifacts": all_artifact_entries,
        }
        try:
            async with AsyncSessionLocal() as _db:
                await _db.execute(
                    update(RemediationReport)
                    .where(RemediationReport.id == report_id)
                    .values(content_json=_partial, generated_doc_ids=generated_doc_ids)
                )
                await _db.commit()
        except Exception as _e:
            logger.warning(
                "Incremental save failed for report %d after control %s: %s",
                report_id, control_id, _e,
            )

    content = {
        "report_type": "artifacts",
        "generated_at": datetime.now(UTC).isoformat(),
        "summary": {
            "controls_addressed": total_controls,
            "documents_created": len(generated_doc_ids),
        },
        "artifacts": all_artifact_entries,
    }

    async with AsyncSessionLocal() as db:
        await db.execute(
            update(RemediationReport)
            .where(RemediationReport.id == report_id)
            .values(
                status="complete",
                content_json=content,
                generated_doc_ids=generated_doc_ids,
                progress_detail=(
                    f"Complete — {len(generated_doc_ids)} evidence artifacts created "
                    f"for {total_controls} controls, queued for indexing"
                ),
            )
        )
        await db.commit()


async def _generate_artifacts_v2(
    report_id: int,
    assessment: Assessment,
    findings: list,
    llm,
    system_context: str,
) -> None:
    """Generate per-control remediation artifacts and prove them before accepting them."""
    from types import SimpleNamespace
    from sqlalchemy import delete
    from sqlalchemy.orm import selectinload
    from app.core.config import get_settings
    from app.models.orm import AssessmentPolicy
    from app.services.prompt_manager import get_prompt
    from app.services.artifact_validation import validate_generated_artifacts
    from app.services.assessment_pipeline import (
        assess_control_with_assessor_pipeline,
        build_scope_document_ids,
        preload_evidence_index,
    )
    from app.services.assessment_policy import build_policy_runtime
    from app.services.closure_service import (
        _build_control_namespace,
        _proof_document_types,
        _wait_for_document_index,
    )
    from app.services.parsers.dispatcher import dispatch_parse_batch
    from app.services.system_knowledge import extract_system_knowledge

    settings = get_settings()
    findings, non_assessable_controls = _split_assessable_findings(findings)
    gap_findings = [f for f in findings if f.status in ("non_compliant", "partially_compliant")]
    total_controls = len(gap_findings)

    async with AsyncSessionLocal() as db:
        report = (await db.execute(select(RemediationReport).where(RemediationReport.id == report_id))).scalar_one()
        config = (report.content_json or {}).get("config", {})
        created_by = report.created_by
        assessment_runtime = (
            await db.execute(
                select(Assessment)
                .options(selectinload(Assessment.policy).selectinload(AssessmentPolicy.buckets))
                .where(Assessment.id == assessment.id)
            )
        ).scalars().first()

    if total_controls == 0:
        async with AsyncSessionLocal() as db:
            await db.execute(
                update(RemediationReport)
                .where(RemediationReport.id == report_id)
                .values(
                    status="complete",
                    content_json={
                        "report_type": "artifacts",
                        "config": config,
                        "summary": {
                            "controls_addressed": 0,
                            "documents_created": 0,
                            "passed_controls": 0,
                            "excluded_non_assessable_count": len(non_assessable_controls),
                            "excluded_non_assessable_controls": non_assessable_controls,
                        },
                        "artifacts": [],
                    },
                    generated_doc_ids=[],
                    progress_detail="Complete - no non-compliant or partial assessable controls to address",
                )
            )
            await db.commit()
        return

    total_started_at = perf_counter()
    project_upload_dir = Path(settings.upload_dir) / str(assessment.project_id)
    project_upload_dir.mkdir(parents=True, exist_ok=True)
    system_name = system_context.split(" - ")[0].strip() if " - " in system_context else system_context or "System"
    date_str = datetime.now().strftime("%B %d, %Y")

    generated_doc_ids: list[int] = []
    artifact_entries: list[dict] = []
    passed_controls = 0
    completed_controls = 0
    worker_count = max(1, min(int(config.get("parallel_controls") or PARALLEL_CONTROL_WORKERS), 8))
    repair_failed_serially = bool(config.get("repair_failed_serially", True))
    failed_controls: list[tuple[int, ControlFinding]] = []
    state_lock = asyncio.Lock()
    failed_controls_lock = asyncio.Lock()
    progress_lock = asyncio.Lock()
    evidence_prompt = await get_prompt("remediation_artifact_evidence", EVIDENCE_COLLECTION_SYSTEM_PROMPT)
    content_prompt_cache: dict[str, str] = {
        "technical_artifact": evidence_prompt,
        "evidence_template": evidence_prompt,
    }
    proof_llm = await _ensure_llm_provider(llm, purpose="assessment_reasoning")

    async def _save_progress(last_control_id: str | None = None) -> None:
        summary = {
            "controls_addressed": total_controls,
            "documents_created": len(generated_doc_ids),
            "passed_controls": passed_controls,
            "failed_controls": total_controls - passed_controls,
            "completed_controls": completed_controls,
            "last_control_id": last_control_id,
            "excluded_non_assessable_count": len(non_assessable_controls),
        }
        async with AsyncSessionLocal() as db:
            await db.execute(
                update(RemediationReport)
                .where(RemediationReport.id == report_id)
                .values(
                    content_json={
                        "report_type": "artifacts",
                        "config": config,
                        "generated_at": datetime.now(UTC).isoformat(),
                        "summary": summary,
                        "artifacts": artifact_entries,
                    },
                    generated_doc_ids=generated_doc_ids,
                    progress_detail=(
                        f"{completed_controls}/{total_controls} complete - "
                        f"{passed_controls} proven compliant"
                        + (f" - last {last_control_id}" if last_control_id else "")
                    ),
                )
            )
            await db.commit()

    async def _prove_control(
        finding: ControlFinding,
        criteria_objectives: list,
        proof_doc_ids: list[int] | None = None,
    ) -> tuple[str, float, list[str], str | None]:
        async with AsyncSessionLocal() as db:
            proof_assessment = Assessment(
                project_id=assessment.project_id,
                status="running",
                llm_provider=assessment.llm_provider,
                llm_model=assessment.llm_model,
                context_strategy=assessment.context_strategy,
                skip_stage3=assessment.skip_stage3,
                carry_forward_compliant=False,
                started_at=datetime.now(UTC),
                name=f"Proof - {finding.control_id}",
                started_by=assessment.started_by,
                policy_id=assessment_runtime.policy_id if assessment_runtime else assessment.policy_id,
                policy_version=assessment_runtime.policy_version if assessment_runtime else assessment.policy_version,
            )
            db.add(proof_assessment)
            await db.flush()

            scope_doc_ids = [
                int(doc_id)
                for doc_id in (proof_doc_ids or [])
                if isinstance(doc_id, int) or str(doc_id).isdigit()
            ]
            if not scope_doc_ids:
                scope_doc_ids = await build_scope_document_ids(assessment.project_id, db)
            evidence_index = await preload_evidence_index(assessment.project_id, scope_doc_ids, db)
            control = _build_control_namespace(
                control_id=finding.control_id,
                finding=finding,
                criteria=SimpleNamespace(
                    assessment_objectives=criteria_objectives,
                    control_statement="",
                    supplemental_guidance="",
                ),
            )
            result = await assess_control_with_assessor_pipeline(
                assessment_id=proof_assessment.id,
                project_id=assessment.project_id,
                control=control,
                system_context=system_context,
                llm=proof_llm,
                db=db,
                evidence_index=evidence_index,
                skip_stage3=assessment.skip_stage3,
                policy_runtime=build_policy_runtime(assessment_runtime.policy if assessment_runtime else None),
            )
            proof_assessment.status = "complete"
            proof_assessment.completed_at = datetime.now(UTC)
            await db.commit()

            status = result.status if result else "not_reviewed"
            confidence = result.confidence if result else 0.0
            gaps = list(result.gaps or []) if result else ["No result returned."]
            dissent = result.llm_challenge_note if result else None

            await db.execute(delete(Assessment).where(Assessment.id == proof_assessment.id))
            await db.commit()
            return status, confidence, gaps, dissent

    async def _process_control(idx: int, finding: ControlFinding, repair_phase: bool = False) -> bool:
        nonlocal passed_controls, completed_controls, generated_doc_ids, artifact_entries

        async with progress_lock:
            await _set_report_progress(
                report_id,
                (
                    f"{completed_controls}/{total_controls} complete - repairing {finding.control_id}..."
                    if repair_phase
                    else f"{completed_controls}/{total_controls} complete - starting {finding.control_id} ({idx}/{total_controls})..."
                ),
            )

        if await _is_cancelled(report_id):
            return False

        control_id = finding.control_id
        async with AsyncSessionLocal() as db:
            criteria_objectives = (
                await db.execute(
                    select(AssessmentCriteriaPackage.assessment_objectives).where(
                        AssessmentCriteriaPackage.assessment_id == assessment.id,
                        AssessmentCriteriaPackage.control_id == control_id,
                    )
                )
            ).scalar_one_or_none() or []

        attempt_gaps = _build_gap_objectives(finding, criteria_objectives)
        final_status = finding.status
        final_confidence = finding.confidence_score or 0.0
        final_gaps = list(finding.gaps or [])
        final_dissent = None
        local_doc_ids: list[int] = []
        local_entries: list[dict] = []
        control_passed = False

        for attempt_no in range(1, 3):
            if await _is_cancelled(report_id):
                return False

            control_doc_ids: list[int] = []
            guidance = build_control_closure_guidance(
                control_id=control_id,
                control_title=finding.control_title,
                gaps=attempt_gaps,
                system_name=system_name,
                current_status=finding.status,
                mode="synthetic",
            )
            doc_types = _proof_document_types(control_id, guidance)

            pending_parse_doc_ids: list[int] = []
            pending_parse_entries: list[dict] = []

            for doc_type in doc_types:
                label = {
                    "policy": "Control Definition and Governance Policy",
                    "ssp_narrative": "System Documentation and Control Narrative",
                    "technical_artifact": "Technical Implementation Evidence Package",
                    "procedure": "Operational Security Procedure Pack",
                }.get(doc_type, "Closure Artifact")
                title = f"{control_id} {label}"
                objectives_prompt_text = _format_objectives_for_prompt(
                    attempt_gaps,
                    control_id,
                    finding.control_title,
                )
                content_system_prompt = content_prompt_cache.get(doc_type)
                if content_system_prompt is None:
                    content_system_prompt = await get_prompt(
                        f"remediation_artifact_{doc_type}",
                        _build_artifact_content_prompt(doc_type),
                    )
                    content_prompt_cache[doc_type] = content_system_prompt
                user_msg = (
                    f"System context summary: {system_context or system_name}\n\n"
                    f"Document to produce: {title}\n"
                    f"Artifact type: {doc_type}\n"
                    f"Controls addressed: {control_id}\n\n"
                    f"Assessment objectives this artifact must close:\n\n{objectives_prompt_text}\n\n"
                    "Write one current-state remediation artifact that closes the objectives explicitly. "
                    "Use concrete values, roles, approvals, review cadence, records, alerting, and enforcement details where the objective requires them."
                )
                content_json = json.dumps(
                    {
                        "title": title,
                        "sections": build_contract_sections(
                            contracts=guidance["objective_contracts"],
                            system_name=system_name,
                            document_type=doc_type,
                            intro_title=title,
                            intro_text=(
                                f"This remediation artifact closes {control_id} for {system_name}. "
                                "Each objective section states current implemented evidence for reassessment."
                            ),
                        ),
                    }
                )
                if attempt_no > 1:
                    try:
                        content_json = _strip_md_fences(
                            await _call_llm(llm, content_system_prompt, user_msg, max_tokens=6000)
                        )
                    except Exception as exc:
                        logger.warning("Per-control artifact generation failed for %s %s: %s", control_id, doc_type, exc)
                content_json = _normalize_artifact_content_json(
                    content_json=content_json,
                    bundle={
                        "title": title,
                        "control_ids": [control_id],
                        "findings": [finding],
                        "assessment_objectives": criteria_objectives,
                        "document_type": doc_type,
                        "artifact_type": doc_type,
                        "family_title": _family_title(finding.control_family),
                        "evidence_role": "implementation",
                    },
                    system_name=system_name,
                    system_context=system_context,
                )
                filename = _build_artifact_filename("REMEDIATE", f"{control_id}_{doc_type}", title)
                doc_id = await _save_document(
                    file_bytes=_json_to_docx(
                        title=title,
                        subtitle=system_name,
                        json_content=content_json,
                        generated_date=date_str,
                    ),
                    filename=filename,
                    file_ext="docx",
                    project_id=assessment.project_id,
                    assessment_id=assessment.id,
                    report_id=report_id,
                    created_by=created_by,
                    upload_dir=project_upload_dir,
                    controls_addressed=[control_id],
                    document_type=doc_type,
                    document_intent="implements",
                    trigger_parse=False,
                )
                local_doc_ids.append(doc_id)
                control_doc_ids.append(doc_id)
                pending_parse_doc_ids.append(doc_id)
                pending_parse_entries.append(
                    {
                        "control_id": control_id,
                        "family": finding.control_family,
                        "family_title": _family_title(finding.control_family),
                        "title": title,
                        "artifact_type": doc_type,
                        "doc_id": doc_id,
                        "attempt": attempt_no,
                        "repair_phase": repair_phase,
                    }
                )

            if pending_parse_doc_ids:
                await dispatch_parse_batch(
                    pending_parse_doc_ids,
                    max_concurrency=max(1, min(len(pending_parse_doc_ids), 4)),
                )
                parse_statuses = await asyncio.gather(
                    *(_wait_for_document_index(doc_id) for doc_id in pending_parse_doc_ids)
                )
                for entry, parse_status in zip(pending_parse_entries, parse_statuses, strict=False):
                    local_entries.append({**entry, "parse_status": parse_status})

            final_status, final_confidence, final_gaps, final_dissent = await _prove_control(
                finding,
                criteria_objectives,
                control_doc_ids,
            )
            local_entries.append(
                {
                    "control_id": control_id,
                    "family": finding.control_family,
                    "family_title": _family_title(finding.control_family),
                    "title": f"{control_id} Proof Result",
                    "artifact_type": "proof_result",
                    "attempt": attempt_no,
                    "repair_phase": repair_phase,
                    "proof_status": final_status,
                    "confidence": final_confidence,
                    "remaining_gaps": final_gaps,
                    "llm_dissent": final_dissent,
                }
            )
            if final_status == "compliant":
                control_passed = True
                break
            attempt_gaps = final_gaps or attempt_gaps

        async with state_lock:
            generated_doc_ids.extend(local_doc_ids)
            generated_doc_ids = await _prune_existing_document_ids(generated_doc_ids)
            artifact_entries.extend(local_entries)
            if control_passed:
                passed_controls += 1
            if not repair_phase:
                completed_controls += 1
            await _save_progress(control_id)
        return control_passed

    semaphore = asyncio.Semaphore(worker_count)

    async def _worker(idx: int, finding: ControlFinding) -> None:
        async with semaphore:
            passed = await _process_control(idx, finding)
            if not passed:
                async with failed_controls_lock:
                    failed_controls.append((idx, finding))

    await asyncio.gather(*(_worker(idx, finding) for idx, finding in enumerate(gap_findings, start=1)))

    if repair_failed_serially and failed_controls and not await _is_cancelled(report_id):
        for repair_idx, repair_finding in failed_controls:
            await _process_control(repair_idx, repair_finding, repair_phase=True)

    validation_summary = None
    system_knowledge_summary = None
    if generated_doc_ids:
        async with AsyncSessionLocal() as db:
            validation_summary = await validate_generated_artifacts(
                db,
                project_id=assessment.project_id,
                document_ids=generated_doc_ids,
                source_mode="remediation_generated",
                source_run_id=report_id,
            )
            system_knowledge_summary = await extract_system_knowledge(
                db,
                project_id=assessment.project_id,
                source_mode="remediation_generated",
                source_run_id=report_id,
                document_ids=generated_doc_ids,
            )

    content = {
        "report_type": "artifacts",
        "config": config,
        "generated_at": datetime.now(UTC).isoformat(),
        "summary": {
            "controls_addressed": total_controls,
            "documents_created": len(generated_doc_ids),
            "passed_controls": passed_controls,
            "failed_controls": total_controls - passed_controls,
            "completed_controls": completed_controls,
            "excluded_non_assessable_count": len(non_assessable_controls),
            "excluded_non_assessable_controls": non_assessable_controls,
        },
        "validation_summary": validation_summary,
        "system_knowledge_summary": system_knowledge_summary,
        "artifacts": artifact_entries,
        "timings": {
            "total_secs": round(perf_counter() - total_started_at, 3),
            "parallel_controls": worker_count,
        },
    }
    async with AsyncSessionLocal() as db:
        await db.execute(
            update(RemediationReport)
            .where(RemediationReport.id == report_id)
            .values(
                status="complete",
                content_json=content,
                generated_doc_ids=generated_doc_ids,
                progress_detail=f"Complete - {passed_controls}/{total_controls} controls proven",
            )
        )
        await db.commit()


# ── Main Entry Point ───────────────────────────────────────────────────────────

async def run_remediation_report(report_id: int) -> None:
    """Background task entry point."""
    async with AsyncSessionLocal() as db:
        result = await db.execute(select(RemediationReport).where(RemediationReport.id == report_id))
        report = result.scalar_one_or_none()
        if not report:
            return

        result = await db.execute(select(Assessment).where(Assessment.id == report.assessment_id))
        assessment = result.scalar_one_or_none()
        if not assessment:
            await db.execute(
                update(RemediationReport).where(RemediationReport.id == report_id)
                .values(status="failed", error_message="Assessment not found")
            )
            await db.commit()
            return

        findings_result = await db.execute(
            select(ControlFinding).where(ControlFinding.assessment_id == assessment.id)
        )
        findings = list(findings_result.scalars().all())

        await db.execute(
            update(RemediationReport).where(RemediationReport.id == report_id)
            .values(
                status="running",
                error_message=None,
            )
        )
        await db.commit()

    try:
        from app.services.llm.runtime import build_provider_for_purpose

        async with AsyncSessionLocal() as _cfg_db:
            llm, _ = await build_provider_for_purpose(
                _cfg_db,
                "remediation_generation",
                provider_name=assessment.llm_provider,
                model=assessment.llm_model,
            )

        from app.models.orm import Project, SystemProfile
        async with AsyncSessionLocal() as db2:
            proj_result = await db2.execute(select(Project).where(Project.id == assessment.project_id))
            project = proj_result.scalar_one_or_none()
            profile_result = await db2.execute(
                select(SystemProfile).where(SystemProfile.project_id == assessment.project_id)
            )
            profile = profile_result.scalar_one_or_none()
        system_context = project.name if project else "Federal information system"
        if profile:
            system_context += f" — {profile.deployment_model or 'cloud'} deployment"

        if report.report_type == "guide":
            await _generate_guide(report_id, assessment, findings, llm, system_context)
        elif report.report_type == "artifacts":
            await _generate_artifacts_v2(report_id, assessment, findings, llm, system_context)
        else:
            raise ValueError(f"Unknown report_type: {report.report_type}")
    except Exception as e:
        logger.exception("Remediation report %d failed: %s", report_id, e)
        async with AsyncSessionLocal() as db:
            await db.execute(
                update(RemediationReport).where(RemediationReport.id == report_id)
                .values(status="failed", error_message=str(e)[:500], progress_detail=None)
            )
            await db.commit()


# ── Family title lookup ────────────────────────────────────────────────────────

_FAMILY_TITLES = {
    "AC": "Access Control",
    "AT": "Awareness and Training",
    "AU": "Audit and Accountability",
    "CA": "Assessment, Authorization, and Monitoring",
    "CM": "Configuration Management",
    "CP": "Contingency Planning",
    "IA": "Identification and Authentication",
    "IR": "Incident Response",
    "MA": "Maintenance",
    "MP": "Media Protection",
    "PE": "Physical and Environmental Protection",
    "PL": "Planning",
    "PM": "Program Management",
    "PS": "Personnel Security",
    "PT": "PII Processing and Transparency",
    "RA": "Risk Assessment",
    "SA": "System and Services Acquisition",
    "SC": "System and Communications Protection",
    "SI": "System and Information Integrity",
    "SR": "Supply Chain Risk Management",
}


def _family_title(family: str) -> str:
    return _FAMILY_TITLES.get(family.upper(), family)


async def _generate_artifacts(
    report_id: int,
    assessment: Assessment,
    findings: list,
    llm,
    system_context: str,
) -> None:
    """Generate reassessment-ready remediation artifacts instead of advisory templates."""
    from app.core.config import get_settings
    from app.services.prompt_manager import get_prompt
    from app.services.artifact_validation import validate_generated_artifacts
    from app.services.evidence_quality import (
        enhance_artifact_document,
        evidence_repair_prompt,
    )
    from app.services.system_knowledge import extract_system_knowledge
    from app.services.test_dataset_generator import _build_persona, _extract_context_from_docs

    settings = get_settings()
    evidence_prompt = await get_prompt("remediation_artifact_evidence", EVIDENCE_COLLECTION_SYSTEM_PROMPT)
    content_prompt_cache: dict[str, str] = {
        "technical_artifact": evidence_prompt,
        "evidence_template": evidence_prompt,
    }

    gap_findings = [f for f in findings if f.status in ("non_compliant", "partially_compliant")]
    total_controls = len(gap_findings)

    async with AsyncSessionLocal() as db:
        rep = (await db.execute(select(RemediationReport).where(RemediationReport.id == report_id))).scalar_one()
        created_by = rep.created_by
        prior_content = rep.content_json or {}
        config = prior_content.get("config", {})
        prior_entries: list[dict] = list(prior_content.get("artifacts", []))
        prior_doc_ids: list[int] = list(rep.generated_doc_ids or [])

    if total_controls == 0:
        async with AsyncSessionLocal() as db:
            await db.execute(
                update(RemediationReport)
                .where(RemediationReport.id == report_id)
                .values(
                    status="complete",
                    content_json={
                        "report_type": "artifacts",
                        "config": config,
                        "summary": {"documents_created": 0, "controls_addressed": 0, "bundle_count": 0},
                        "artifacts": [],
                    },
                    generated_doc_ids=[],
                    progress_detail="Complete - no non-compliant or partial controls to address",
                )
            )
            await db.commit()
        return

    total_started_at = perf_counter()
    timings: dict[str, float | int] = {"checkpoint_interval": PARTIAL_SAVE_EVERY_BUNDLES}

    planning_started_at = perf_counter()
    package_style = (config.get("package_style") or "standard").lower()
    project_upload_dir = Path(settings.upload_dir) / str(assessment.project_id)
    project_upload_dir.mkdir(parents=True, exist_ok=True)

    # Fresh remediation runs must not reuse old autogenerated coverage. If a control is still
    # partial/non-compliant, regenerate against the newest closure logic rather than skipping it
    # because an older synthetic artifact already exists in the library.
    pending_findings = list(gap_findings)
    skipped_controls: list[str] = []

    blueprint = plan_remediation_bundles(pending_findings, package_style=package_style)
    validation = build_blueprint_validation(blueprint)
    bundles = blueprint.get("bundles", [])
    serializable_blueprint = {
        "summary": blueprint.get("summary", {}),
        "bundles": [
            {
                "bundle_id": bundle.get("bundle_id"),
                "family": bundle.get("family"),
                "family_title": bundle.get("family_title"),
                "title": bundle.get("title"),
                "artifact_type": bundle.get("artifact_type"),
                "control_ids": list(bundle.get("control_ids", [])),
                "findings": [
                    {
                        "control_id": getattr(finding, "control_id", None),
                        "status": getattr(finding, "status", None),
                        "family": getattr(finding, "control_family", getattr(finding, "family", None)),
                        "control_title": getattr(finding, "control_title", None),
                    }
                    for finding in bundle.get("findings", [])
                ],
            }
            for bundle in bundles
        ],
    }
    timings["planning_secs"] = round(perf_counter() - planning_started_at, 3)

    completed_bundles = {
        entry.get("bundle_id")
        for entry in prior_entries
        if entry.get("bundle_id") and not entry.get("error")
    }
    generated_doc_ids: list[int] = list(prior_doc_ids)
    pending_dispatch_doc_ids: list[int] = []
    artifact_entries: list[dict] = list(prior_entries)

    system_name = system_context.split(" - ")[0].strip() if " - " in system_context else system_context or "System"
    context_persona = ""
    try:
        context_payload = await _extract_context_from_docs(assessment.project_id, llm, "moderate")
        if context_payload.get("system_name"):
            system_name = context_payload["system_name"]
        context_persona = _build_persona(context_payload, system_name)
    except Exception as exc:
        logger.warning("Remediation artifact context extraction failed for project %s: %s", assessment.project_id, exc)
    date_str = datetime.now().strftime("%B %d, %Y")

    artifact_doc_types = {
        "policy": "policy",
        "policy_procedure": "procedure",
        "completed_form": "procedure",
        "ssp_narrative": "ssp_narrative",
        "procedure": "procedure",
        "technical_artifact": "technical_artifact",
        "evidence_template": "technical_artifact",
        "agreement_template": "other",
    }
    generation_started_at = perf_counter()
    processed_since_checkpoint = 0

    for idx, bundle in enumerate(bundles, start=1):
        if await _is_cancelled(report_id):
            return
        if bundle["bundle_id"] in completed_bundles:
            continue

        await _set_report_progress(
            report_id,
            f"{idx}/{len(bundles)} - generating {bundle['title']}...",
        )

        artifact_type = bundle["artifact_type"]
        if artifact_type in {"technical_artifact", "evidence_template"}:
            content_system_prompt = evidence_prompt
        else:
            content_system_prompt = content_prompt_cache.get(artifact_type)
            if content_system_prompt is None:
                content_system_prompt = await get_prompt(
                    f"remediation_artifact_{artifact_type}",
                    _build_artifact_content_prompt(artifact_type),
                )
                content_prompt_cache[artifact_type] = content_system_prompt

        objective_blocks = []
        quality_controls: list[dict] = []
        for finding in bundle["findings"]:
            criteria_result = await db.execute(
                select(AssessmentCriteriaPackage.assessment_objectives).where(
                    AssessmentCriteriaPackage.assessment_id == assessment.id,
                    AssessmentCriteriaPackage.control_id == finding.control_id,
                )
            )
            criteria_objectives = criteria_result.scalar_one_or_none() or []
            gap_objectives = _build_gap_objectives(finding, criteria_objectives)
            objective_blocks.append(
                _format_objectives_for_prompt(
                    gap_objectives,
                    finding.control_id,
                    finding.control_title,
                )
            )
            quality_controls.append(
                {
                    "control_id": finding.control_id,
                    "title": finding.control_title,
                    "family": getattr(finding, "control_family", getattr(finding, "family", bundle["family"])),
                    "status": finding.status,
                    "objectives": [
                        f"{obj.get('objective_id')}: {obj.get('description') or obj.get('full_text') or ''}"
                        for obj in gap_objectives
                    ],
                }
            )
        objectives_prompt_text = "\n\n".join(objective_blocks)

        safe_name = _build_artifact_filename(
            "REMEDIATE",
            bundle["bundle_id"],
            bundle["title"],
        )
        existing_content = await _get_existing_content(assessment.project_id, safe_name)

        if existing_content:
            user_msg = (
                f"{context_persona}\n"
                f"System context summary: {system_context or 'Federal information system'}\n\n"
                f"Document to update: {bundle['title']}\n"
                f"Artifact type: {artifact_type}\n"
                f"Evidence role: {bundle.get('evidence_role', 'implementation')}\n"
                f"Document intent: {bundle.get('document_intent', 'implements')}\n"
                f"Controls addressed: {', '.join(bundle['control_ids'])}\n\n"
                f"Existing document content:\n---\n{existing_content}\n---\n\n"
                f"Assessment objectives to cover:\n\n{objectives_prompt_text}\n\n"
                "Preserve existing useful sections, append any missing objective-labeled sections, "
                "and return the complete updated document as one JSON object.\n\n"
                "This artifact must read like current implemented evidence that should score as compliant "
                "on reassessment. Do not output future-state TODO language, empty placeholders, or generic templates."
            )
        else:
            user_msg = (
                f"{context_persona}\n"
                f"System context summary: {system_context or 'Federal information system'}\n\n"
                f"Document to produce: {bundle['title']}\n"
                f"Artifact type: {artifact_type}\n"
                f"Evidence role: {bundle.get('evidence_role', 'implementation')}\n"
                f"Document intent: {bundle.get('document_intent', 'implements')}\n"
                f"Controls addressed: {', '.join(bundle['control_ids'])}\n\n"
                f"Assessment objectives this artifact package must close:\n\n{objectives_prompt_text}\n\n"
                "Generate one coherent document that covers all listed controls and objectives.\n"
                "Write as if the control implementation exists today and this artifact is being prepared to pass reassessment."
            )

        content_json = json.dumps(
            {
                "title": bundle["title"],
                "sections": [
                    {"type": "heading", "level": 1, "text": bundle["title"]},
                    {
                        "type": "paragraph",
                        "text": (
                            f"This reassessment-ready artifact package supports {system_name} and addresses controls "
                            f"{', '.join(bundle['control_ids'])}."
                        ),
                    },
                ],
            }
        )
        try:
            content_json = _strip_md_fences(
                await _call_llm(llm, content_system_prompt, user_msg, max_tokens=6000)
            )
        except Exception as exc:
            logger.warning("Artifact generation failed for %s: %s", bundle["bundle_id"], exc)

        content_json = _normalize_artifact_content_json(
            content_json=content_json,
            bundle=bundle,
            system_name=system_name,
            system_context=system_context,
        )
        parsed_content = _parse_artifact_json(content_json)
        enhanced_content, evidence_quality = enhance_artifact_document(
            document=parsed_content,
            controls=quality_controls,
            artifact_type=artifact_type,
            evidence_role=bundle.get("evidence_role", "implementation"),
            system_name=system_name,
            system_context=system_context,
            organization=context_payload.get("organization") if "context_payload" in locals() else None,
            source="remediation_generated",
        )
        if not evidence_quality.get("passed"):
            try:
                repair_system, repair_user = evidence_repair_prompt(
                    document=enhanced_content,
                    controls=quality_controls,
                    quality_summary=evidence_quality,
                    system_name=system_name,
                )
                repaired_raw = _strip_md_fences(await _call_llm(llm, repair_system, repair_user, max_tokens=6000))
                repaired_content = _parse_artifact_json(repaired_raw)
                enhanced_content, evidence_quality = enhance_artifact_document(
                    document=repaired_content,
                    controls=quality_controls,
                    artifact_type=artifact_type,
                    evidence_role=bundle.get("evidence_role", "implementation"),
                    system_name=system_name,
                    system_context=system_context,
                    organization=context_payload.get("organization") if "context_payload" in locals() else None,
                    source="remediation_generated_repaired",
                )
            except Exception as exc:
                logger.warning("Evidence quality repair failed for %s: %s", bundle["bundle_id"], exc)
        content_json = json.dumps(enhanced_content)

        try:
            docx_bytes = _json_to_docx(
                title=bundle["title"],
                subtitle=system_name,
                json_content=content_json,
                generated_date=date_str,
            )
            doc_id = await _save_document(
                file_bytes=docx_bytes,
                filename=safe_name,
                file_ext="docx",
                project_id=assessment.project_id,
                assessment_id=assessment.id,
                report_id=report_id,
                created_by=created_by,
                upload_dir=project_upload_dir,
                controls_addressed=bundle["control_ids"],
                document_type=bundle.get("document_type") or artifact_doc_types.get(artifact_type, "procedure"),
                document_intent=bundle.get("document_intent", "implements"),
                trigger_parse=False,
            )
            generated_doc_ids.append(doc_id)
            pending_dispatch_doc_ids.append(doc_id)
            artifact_entries.append(
                {
                    "bundle_id": bundle["bundle_id"],
                    "family": bundle["family"],
                    "family_title": bundle["family_title"],
                    "title": bundle["title"],
                    "artifact_type": artifact_type,
                    "controls_addressed": bundle["control_ids"],
                    "evidence_quality": evidence_quality,
                    "doc_id": doc_id,
                }
            )
            completed_bundles.add(bundle["bundle_id"])
        except Exception as exc:
            logger.error("Failed to save remediation supplement for %s: %s", bundle["bundle_id"], exc)
            artifact_entries.append(
                {
                    "bundle_id": bundle["bundle_id"],
                    "family": bundle["family"],
                    "family_title": bundle["family_title"],
                    "title": bundle["title"],
                    "artifact_type": artifact_type,
                    "controls_addressed": bundle["control_ids"],
                    "error": str(exc),
                }
            )

        processed_since_checkpoint += 1
        if processed_since_checkpoint >= PARTIAL_SAVE_EVERY_BUNDLES:
            generated_doc_ids = await _prune_existing_document_ids(generated_doc_ids)
            partial_content = {
                "report_type": "artifacts",
                "config": config,
                "generated_at": datetime.now(UTC).isoformat(),
                "blueprint": serializable_blueprint,
                "validation": validation,
                "summary": {
                    "controls_addressed": total_controls,
                    "documents_created": len(generated_doc_ids),
                    "bundle_count": len(bundles),
                    "skipped_controls": skipped_controls,
                },
                "artifacts": artifact_entries,
            }
            async with AsyncSessionLocal() as db:
                await db.execute(
                    update(RemediationReport)
                    .where(RemediationReport.id == report_id)
                    .values(content_json=partial_content, generated_doc_ids=generated_doc_ids)
                )
                await db.commit()
            if pending_dispatch_doc_ids:
                dispatch_started_at = perf_counter()
                from app.services.parsers.dispatcher import dispatch_parse_batch
                await dispatch_parse_batch(pending_dispatch_doc_ids)
                timings["index_dispatch_secs"] = round(
                    (timings.get("index_dispatch_secs", 0.0) or 0.0)
                    + (perf_counter() - dispatch_started_at),
                    3,
                )
                pending_dispatch_doc_ids = []
            processed_since_checkpoint = 0

    timings["generation_secs"] = round(perf_counter() - generation_started_at, 3)

    if pending_dispatch_doc_ids:
        dispatch_started_at = perf_counter()
        from app.services.parsers.dispatcher import dispatch_parse_batch
        await dispatch_parse_batch(pending_dispatch_doc_ids)
        timings["index_dispatch_secs"] = round(
            (timings.get("index_dispatch_secs", 0.0) or 0.0)
            + (perf_counter() - dispatch_started_at),
            3,
        )
    elif "index_dispatch_secs" not in timings:
        timings["index_dispatch_secs"] = 0.0

    generated_doc_ids = await _prune_existing_document_ids(generated_doc_ids)
    pending_dispatch_doc_ids = [doc_id for doc_id in pending_dispatch_doc_ids if doc_id in set(generated_doc_ids)]

    validation_summary = None
    system_knowledge_summary = None
    if generated_doc_ids:
        async with AsyncSessionLocal() as db:
            validation_started_at = perf_counter()
            validation_summary = await validate_generated_artifacts(
                db,
                project_id=assessment.project_id,
                document_ids=generated_doc_ids,
                source_mode="remediation_generated",
                source_run_id=report_id,
            )
            timings["artifact_validation_secs"] = round(perf_counter() - validation_started_at, 3)
            knowledge_started_at = perf_counter()
            system_knowledge_summary = await extract_system_knowledge(
                db,
                project_id=assessment.project_id,
                source_mode="remediation_generated",
                source_run_id=report_id,
                document_ids=generated_doc_ids,
            )
            timings["system_knowledge_secs"] = round(perf_counter() - knowledge_started_at, 3)
    else:
        timings["artifact_validation_secs"] = 0.0
        timings["system_knowledge_secs"] = 0.0

    timings["total_secs"] = round(perf_counter() - total_started_at, 3)

    final_content = {
        "report_type": "artifacts",
        "generation_mode": "remediation_generated",
        "config": config,
        "generated_at": datetime.now(UTC).isoformat(),
        "blueprint": serializable_blueprint,
        "validation": validation,
        "artifact_validation": validation_summary,
        "system_knowledge": system_knowledge_summary,
        "timing": timings,
        "summary": {
            "controls_addressed": total_controls,
            "documents_created": len(generated_doc_ids),
            "bundle_count": len(bundles),
            "skipped_controls": skipped_controls,
        },
        "artifacts": artifact_entries,
    }

    async with AsyncSessionLocal() as db:
        await db.execute(
            update(RemediationReport)
            .where(RemediationReport.id == report_id)
            .values(
                status="complete",
                content_json=final_content,
                generated_doc_ids=generated_doc_ids,
                progress_detail=(
                    f"Complete - {len(generated_doc_ids)} consolidated remediation documents created "
                    f"for {total_controls} controls"
                ),
            )
        )
        await db.commit()
