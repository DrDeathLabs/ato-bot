"""
Test Dataset Generator — standalone, project-level ATO evidence package generator.

Generates one targeted document per control in the project's impact baseline.
Before generating, extracts real system context from already-uploaded project
documents (if any) so generated content is consistent with the existing package.

Each generated document:
  • Is written in present-tense "implemented" language (not future-tense SHALL)
  • Contains a dedicated labeled section per NIST 800-53A assessment objective
    so the assessment engine can retrieve and score each objective independently
  • Is tagged with the correct document_type (policy | procedure |
    technical_artifact | ssp_narrative) so the Stage 2 LLM applies correct weighting
  • Has document_intent = "implements" so it is treated as current-state evidence

Document types produced by control class:
  policy             — X-1 family policy controls
  procedure          — operational process controls
  technical_artifact — configuration / infrastructure controls
  ssp_narrative      — PL, CA, RA planning controls
"""
from __future__ import annotations

import asyncio
import io
import json
import logging
import re
from datetime import UTC, datetime
from pathlib import Path
from time import perf_counter
from uuid import uuid4

from sqlalchemy import select, update

from app.core.database import AsyncSessionLocal
from app.models.orm import Assessment, Document, Project, RemediationReport, SystemProfile, TestDatasetJob
from app.services.closure_guidance import (
    build_control_closure_guidance,
    build_contract_sections,
    format_contracts_for_prompt,
    sections_satisfy_contracts,
)
from app.services.evidence_view import build_system_context_from_evidence
from app.services.implementation_statements import (
    build_control_statement_generation_guidance,
    synthesize_control_implementation_statement,
)
from app.services.package_generation import (
    build_blueprint_validation,
    build_expected_outcomes,
    plan_test_dataset_bundles,
)

logger = logging.getLogger(__name__)

PARTIAL_SAVE_EVERY_BUNDLES = 3
CONTEXT_EXTRACTION_TIMEOUT_SECS = 45
BUNDLE_GENERATION_TIMEOUT_SECS = 90
BUNDLE_REPAIR_TIMEOUT_SECS = 60


async def _complete_with_timeout(
    llm,
    system_prompt: str,
    user_prompt: str,
    *,
    timeout_secs: int,
    label: str,
) -> str:
    """Keep one slow provider call from stalling an entire package run."""
    try:
        return await asyncio.wait_for(
            llm.complete(system_prompt, user_prompt),
            timeout=timeout_secs,
        )
    except asyncio.TimeoutError as exc:
        raise TimeoutError(f"{label} timed out after {timeout_secs}s") from exc

# ── Control classification ────────────────────────────────────────────────────

TECHNICAL_CONTROL_IDS: frozenset[str] = frozenset({
    "SC-2", "SC-3", "SC-4", "SC-5", "SC-7", "SC-7(4)", "SC-7(7)",
    "SC-10", "SC-15", "SC-18", "SC-20", "SC-22", "SC-28", "SC-28(1)", "SC-39",
    "SI-3", "SI-3(1)", "SI-3(2)", "SI-4", "SI-4(4)", "SI-6", "SI-7", "SI-8", "SI-8(1)",
    "AU-3(1)", "AU-4", "AU-7(1)", "AU-9(4)",
    "CP-6", "CP-6(1)", "CP-6(3)", "CP-7", "CP-7(1)", "CP-7(2)", "CP-7(3)",
    "CP-8", "CP-8(1)", "CA-7(1)", "SA-17", "SC-13", "SC-17",
    "AC-17", "AC-18", "AC-19", "IA-7",
})

SSP_FAMILIES: frozenset[str] = frozenset({"PL", "CA", "RA"})


def _classify_control(control_id: str) -> tuple[str, str, str]:
    """Return (doc_type_label, document_type_field, content_mode)."""
    family = control_id.split("-")[0].upper()
    base = re.sub(r"\(.*\)", "", control_id).strip()

    if control_id in TECHNICAL_CONTROL_IDS:
        return ("Technical Implementation Guide", "technical_artifact", "technical")
    if family in SSP_FAMILIES:
        return ("SSP Control Implementation Statement", "ssp_narrative", "ssp")
    if family == "AT":
        return ("Security Awareness and Training Record", "procedure", "training")
    if base.endswith("-1"):
        return ("Security Policy", "policy", "policy")
    if family in ("PE", "PS", "MA", "MP"):
        return ("Security Procedure", "procedure", "procedure")
    if family == "CP":
        return ("Contingency Planning Procedure", "procedure", "procedure")
    if "(" in control_id:
        return ("Security Procedure", "procedure", "procedure")
    return ("Security Procedure and Policy", "procedure", "procedure")


# ── Context extraction from existing documents ────────────────────────────────

CONTEXT_EXTRACTION_PROMPT = """\
You are reading excerpts from an existing ATO documentation package.
Extract the following information from these excerpts and return ONLY a JSON object:
{
  "system_name": "exact system/application name",
  "organization": "full organization or contractor name",
  "impact_level": "Low | Moderate | High",
  "cloud_platform": "e.g. AWS GovCloud, Azure Government, on-premise",
  "app_stack": "key technologies (brief list)",
  "security_tools": "key security tools mentioned",
  "system_owner": "name and title of system owner",
  "isso": "name and title of ISSO",
  "ao": "name and title of Authorizing Official",
  "mission": "one sentence: what this system does"
}

If any field is not found in the excerpts, use null.
Return ONLY the JSON object — no commentary."""

FALLBACK_PERSONA = {
    "system_name": None,
    "organization": "Keystone Federal Solutions, LLC",
    "impact_level": None,
    "cloud_platform": "AWS GovCloud (US-East), FedRAMP authorized",
    "app_stack": "Python/Django REST API, PostgreSQL 14, Redis, Node.js frontend",
    "security_tools": (
        "Tenable.io (vulnerability scanning), CrowdStrike Falcon (EDR), "
        "Splunk Cloud SIEM, AWS CloudTrail, AWS Config, Qualys, ServiceNow"
    ),
    "system_owner": "Michael Okonkwo, Senior Vice President of Technology",
    "isso": "Priya Venkataraman, ISSO",
    "ao": "Sarah J. Blackwood, Deputy CISO, Office of Management Review",
    "mission": "Federal program data management and analytics platform",
}


async def _extract_context_from_docs(
    project_id: int, llm, impact_level: str
) -> dict:
    """
    Pull text from already-indexed project documents and use the LLM to extract
    system context. Falls back to the default Keystone Federal persona if no
    documents are indexed or extraction fails.
    """
    async with AsyncSessionLocal() as db:
        context_text = await build_system_context_from_evidence(project_id, db, max_chars=6000)

    if not context_text or context_text == "No system context documents available.":
        logger.info("No indexed documents found for project %d — using fallback persona", project_id)
        return {**FALLBACK_PERSONA, "impact_level": impact_level.capitalize()}

    try:
        raw = await _complete_with_timeout(
            llm,
            CONTEXT_EXTRACTION_PROMPT,
            f"Document excerpts:\n\n{context_text}",
            timeout_secs=CONTEXT_EXTRACTION_TIMEOUT_SECS,
            label="Context extraction",
        )
        start = raw.find("{")
        end = raw.rfind("}") + 1
        if start != -1 and end > start:
            extracted = json.loads(raw[start:end])
            # Merge with fallback for any null fields
            merged = {**FALLBACK_PERSONA}
            for k, v in extracted.items():
                if v and str(v).strip() not in ("null", "None", ""):
                    merged[k] = v
            if not merged.get("impact_level"):
                merged["impact_level"] = impact_level.capitalize()
            logger.info("Extracted system context for project %d: %s", project_id, merged.get("system_name"))
            return merged
    except Exception as e:
        logger.warning("Context extraction failed for project %d: %s", project_id, e)

    return {**FALLBACK_PERSONA, "impact_level": impact_level.capitalize()}


def _build_persona(ctx: dict, system_name_override: str) -> str:
    """Build the system persona block injected into every generation prompt."""
    system_name = ctx.get("system_name") or system_name_override
    return f"""\
You are writing an official, production-grade federal security document.
Use ALL of the following system details throughout the document to make it specific and credible.
Do NOT invent different organization or tool names — use exactly what is listed here.

SYSTEM CONTEXT:
  System Name: {system_name}
  Organization: {ctx.get("organization", "Keystone Federal Solutions, LLC")}
  Impact Level: {ctx.get("impact_level", "Moderate")} (FIPS 199)
  Cloud Platform: {ctx.get("cloud_platform", "AWS GovCloud (US-East)")}
  Application Stack: {ctx.get("app_stack", "Python/Django, PostgreSQL, Redis, Node.js")}
  Security Tools: {ctx.get("security_tools", "Tenable.io, CrowdStrike Falcon, Splunk SIEM, AWS CloudTrail, AWS Config, ServiceNow")}
  System Owner: {ctx.get("system_owner", "Michael Okonkwo, SVP Technology")}
  ISSO: {ctx.get("isso", "Priya Venkataraman, ISSO")}
  Authorizing Official: {ctx.get("ao", "Sarah J. Blackwood, Deputy CISO")}
  Mission: {ctx.get("mission", "Federal program data management and analytics")}
  Assessment Period: FY2025 Annual Authorization Review
"""


# ── Content prompts ───────────────────────────────────────────────────────────

_OUTPUT_FORMAT = """\

OUTPUT FORMAT — return ONLY a valid JSON object, no markdown code fences:
{
  "title": "Full document title including system name",
  "sections": [
    {"type": "heading", "level": 1, "text": "Section heading"},
    {"type": "paragraph", "text": "Plain prose — no markdown symbols"},
    {"type": "numbered_list", "items": ["Step text"]},
    {"type": "bullet_list", "items": ["Item text"]},
    {"type": "table", "headers": ["Col1", "Col2"], "rows": [["val", "val"]]}
  ]
}
Return ONLY the JSON object."""


def _build_prompt(content_mode: str, persona: str, system_name: str) -> str:
    if content_mode == "technical":
        return f"""\
{persona}
You are writing a TECHNICAL IMPLEMENTATION GUIDE for {system_name}.
This will be indexed by an AI assessor retrieving chunks by vector similarity to NIST 800-53A objectives.

Before the per-objective sections, add a heading named "Control Implementation Summary"
followed by 2-4 human-readable paragraphs that synthesize the full control as one coherent
implementation statement.

Write in PRESENT TENSE describing CURRENT implementation. Not future-tense requirements.
"is configured", "is enabled", "are enrolled", "has been implemented" — not "shall be".

MANDATORY STRUCTURE — for EVERY assessment objective:
  Heading: "[OBJECTIVE_ID] — Technical Implementation"
  First sentence: "This section satisfies NIST 800-53A assessment objective [OBJECTIVE_ID]."
  Body: specific technical description naming AWS services, tool names, config settings,
    scan schedules, key ARNs/IDs (realistic but fictitious), who verified it and when.

Include a Verification Record table at the end:
  Columns: Objective ID | Verification Method | Last Verified | Result
{_OUTPUT_FORMAT}"""

    if content_mode == "policy":
        return f"""\
{persona}
You are writing a SECURITY POLICY for {system_name}.
This will be indexed by an AI assessor retrieving chunks by vector similarity to NIST 800-53A objectives.

Before the per-objective sections, add a heading named "Control Implementation Summary"
followed by 2-4 human-readable paragraphs that synthesize the full control as one coherent
implementation statement.

MANDATORY STRUCTURE — for EVERY assessment objective:
  Heading: "[OBJECTIVE_ID] — [Short Policy Title]"
  First sentence: "This section satisfies NIST 800-53A assessment objective [OBJECTIVE_ID]."
  Body: specific policy statement. Name the organization, roles (ISSO, System Owner),
    specific timelines and thresholds. Present tense: "maintains", "reviews", "requires".

Required document sections: Purpose, Scope, Authority, [per-objective sections],
  Roles and Responsibilities (table), Enforcement, Review Schedule, Document Control.
{_OUTPUT_FORMAT}"""

    if content_mode == "ssp":
        return f"""\
{persona}
You are writing an SSP SECTION 9 IMPLEMENTATION STATEMENT for {system_name}.
This will be indexed by an AI assessor retrieving chunks by vector similarity to NIST 800-53A objectives.

Write as the ISSO describing current implementation. SSP narrative format.

Before the per-objective sections, add a heading named "Control Implementation Summary"
followed by 2-4 human-readable paragraphs that synthesize the full control as one coherent
implementation statement.

MANDATORY STRUCTURE — for EVERY assessment objective:
  Heading: "[OBJECTIVE_ID] — [Short Title]"
  Paragraph 1: "This section satisfies NIST 800-53A assessment objective [OBJECTIVE_ID]."
  Paragraph 2: "Implementation: [how this is currently met — specific tools, roles, processes]"
  Paragraph 3: "Evidence: [specific document, ServiceNow ticket, AWS Config rule, scan result]"
  Paragraph 4: "Responsible Party: [role and name]"

End with a summary table: Objective ID | Status | Evidence Reference | Owner
{_OUTPUT_FORMAT}"""

    if content_mode == "training":
        return f"""\
{persona}
You are writing a SECURITY AWARENESS AND TRAINING RECORD for {system_name}.
This will be indexed by an AI assessor.

Before the per-objective sections, add a heading named "Control Implementation Summary"
followed by 2-4 human-readable paragraphs that synthesize the full control as one coherent
implementation statement.

MANDATORY STRUCTURE — for EVERY assessment objective:
  Heading: "[OBJECTIVE_ID] — [Training Requirement Title]"
  First sentence: "This section satisfies NIST 800-53A assessment objective [OBJECTIVE_ID]."
  Body: describe the training program specifics — platform (KnowBe4), completion tracking
    (ServiceNow), frequency, completion rate, ISSO oversight.

Include a TRAINING COMPLETION TABLE with 12 representative rows:
  Columns: Employee ID | Role | Module | Completion Date | Score | Status

Required sections: Training Program Overview, [per-objective sections],
  Completion Records (table), Annual Review Summary, Document Control.
{_OUTPUT_FORMAT}"""

    # Default: procedure
    return f"""\
{persona}
You are writing an operational SECURITY PROCEDURE for {system_name}.
This will be indexed by an AI assessor retrieving chunks by vector similarity to NIST 800-53A objectives.

Write in present-tense operational language describing HOW this is currently carried out.
Name specific tools (ServiceNow, CrowdStrike, Tenable.io, AWS Console), roles, and cadence.

Before the per-objective sections, add a heading named "Control Implementation Summary"
followed by 2-4 human-readable paragraphs that synthesize the full control as one coherent
implementation statement.

MANDATORY STRUCTURE — for EVERY assessment objective:
  Heading: "[OBJECTIVE_ID] — [Short Procedure Title]"
  First sentence: "This section satisfies NIST 800-53A assessment objective [OBJECTIVE_ID]."
  Body: step-by-step procedure steps, tool references, responsible roles, frequency.

Required sections: Purpose and Scope, Prerequisites, [per-objective sections],
  Verification Steps, Record-Keeping (ServiceNow, retention period),
  Roles and Responsibilities, Document Control.
{_OUTPUT_FORMAT}"""


# ── Objectives helper ─────────────────────────────────────────────────────────

def _objectives_from_control(control) -> list[dict]:
    objectives = []
    raw = getattr(control, "assessment_objectives", None) or []
    if isinstance(raw, list):
        for obj in raw:
            if isinstance(obj, dict):
                oid = obj.get("id") or obj.get("objective_id") or control.display_id
                desc = obj.get("description") or obj.get("prose") or obj.get("text") or ""
                objectives.append({"objective_id": oid, "description": desc})
            elif isinstance(obj, str):
                objectives.append({"objective_id": control.display_id, "description": obj})

    if not objectives:
        stmt = getattr(control, "statement", "") or ""
        for letter, text in re.findall(r'\(([a-z])\)\s+([^(]+)', stmt):
            objectives.append({
                "objective_id": f"{control.display_id}{letter}",
                "description": text.strip()[:200],
            })

    if not objectives:
        objectives.append({
            "objective_id": control.display_id,
            "description": getattr(control, "title", control.display_id),
        })
    return objectives


def _format_objectives(objectives: list[dict], control_id: str, control_title: str) -> str:
    contracts = build_control_closure_guidance(
        control_id=control_id,
        control_title=control_title,
        gaps=objectives,
        system_name="the system",
        mode="synthetic",
    )["objective_contracts"]
    lines = [
        f"Control: {control_id} — {control_title}",
        "",
        "ASSESSMENT OBJECTIVES — each MUST get a dedicated labeled section:",
        "",
    ]
    for obj in objectives:
        oid = obj["objective_id"]
        lines += [
            f"  OBJECTIVE ID: {oid}",
            f"  REQUIREMENT: {obj['description']}",
            f"  REQUIRED HEADING: \"{oid} — [descriptive title]\"",
            f"  REQUIRED FIRST SENTENCE: \"This section satisfies NIST 800-53A assessment objective {oid}.\"",
            "",
        ]
    lines.append(format_contracts_for_prompt(contracts))
    return "\n".join(lines)


# ── Word document builder ─────────────────────────────────────────────────────

def _build_docx(title: str, sections: list[dict], org: str) -> bytes:
    from docx import Document as DocxDoc
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDoc()
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run("Organization: ").bold = True
    meta.add_run(org)
    meta.add_run("    |    Classification: ").bold = True
    meta.add_run("Controlled Unclassified Information (CUI)")
    meta.add_run("    |    Generated: ").bold = True
    meta.add_run(datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph()

    for s in sections:
        t = s.get("type", "paragraph")
        if t == "heading":
            doc.add_heading(s.get("text", ""), min(int(s.get("level", 1)), 4))
        elif t == "paragraph":
            if s.get("text"):
                doc.add_paragraph(s["text"])
        elif t == "numbered_list":
            for item in s.get("items") or []:
                doc.add_paragraph(item, style="List Number")
        elif t == "bullet_list":
            for item in s.get("items") or []:
                doc.add_paragraph(item, style="List Bullet")
        elif t == "table":
            headers = s.get("headers", [])
            rows = s.get("rows", [])
            if headers:
                tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
                tbl.style = "Table Grid"
                for i, h in enumerate(headers):
                    cell = tbl.rows[0].cells[i]
                    cell.text = h
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                for ri, row_data in enumerate(rows):
                    for ci, val in enumerate(row_data):
                        if ci < len(tbl.rows[ri + 1].cells):
                            tbl.rows[ri + 1].cells[ci].text = str(val)
                doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _safe_filename_component(text: str, max_length: int = 80) -> str:
    cleaned = re.sub(r"[^\w\s-]", "", text).strip()
    cleaned = re.sub(r"\s+", "_", cleaned)
    return cleaned[:max_length]


def _build_generated_filename(prefix: str, unique_key: str, title: str) -> str:
    parts = [
        _safe_filename_component(prefix, 24),
        _safe_filename_component(unique_key, 64),
        _safe_filename_component(title, 72),
    ]
    return f"{'_'.join(part for part in parts if part)}.docx"


# ── Save document record ──────────────────────────────────────────────────────

async def _save_doc(
    file_bytes: bytes,
    filename: str,
    project_id: int,
    upload_dir: Path,
    created_by: int,
    control_id: str,
    document_type: str,
    document_intent: str = "implements",
    controls_addressed: list[str] | None = None,
    source_assessment_id: int | None = None,
    source_remediation_report_id: int | None = None,
    trigger_parse: bool = True,
) -> int:
    """Write bytes to disk and create a Document record. Return doc id."""
    import asyncio as _asyncio
    import hashlib

    file_hash = hashlib.sha256(file_bytes).hexdigest()
    file_path = upload_dir / f"{uuid4()}.docx"

    async with AsyncSessionLocal() as db:
        # Dedup: remove old autogenerated doc with same filename for this project
        old_result = await db.execute(
            select(Document).where(
                Document.project_id == project_id,
                Document.filename == filename,
                Document.autogenerated == True,  # noqa: E712
            )
        )
        for old in old_result.scalars().all():
            try:
                Path(old.file_path).unlink(missing_ok=True)
            except Exception:
                pass
            await db.delete(old)
        await db.flush()

        file_path.write_bytes(file_bytes)

        doc = Document(
            project_id=project_id,
            filename=filename,
            file_path=str(file_path),
            file_type="docx",
            file_hash=file_hash,
            file_size_bytes=len(file_bytes),
            parse_status="pending",
            uploaded_by=created_by,
            autogenerated=True,
            artifact_status="draft",
            evidence_eligible=False,
            source_assessment_id=source_assessment_id,
            source_remediation_report_id=source_remediation_report_id,
            artifact_controls=controls_addressed or [control_id],
            document_type=document_type,
            document_intent=document_intent,
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)
        doc_id = doc.id

    if trigger_parse:
        from app.services.parsers.dispatcher import dispatch_parse
        _asyncio.create_task(dispatch_parse(doc_id))
    return doc_id


# ── Progress / cancel helpers ─────────────────────────────────────────────────

async def _set_progress(job_id: int, detail: str) -> None:
    async with AsyncSessionLocal() as db:
        await db.execute(
            update(TestDatasetJob)
            .where(TestDatasetJob.id == job_id)
            .values(progress_detail=detail)
        )
        await db.commit()


async def _is_cancelled(job_id: int) -> bool:
    async with AsyncSessionLocal() as db:
        result = await db.execute(
            select(TestDatasetJob.status).where(TestDatasetJob.id == job_id)
        )
        row = result.one_or_none()
        return row is None or row[0] not in ("running", "pending")


async def _set_run_progress(run_type: str, run_id: int, detail: str) -> None:
    async with AsyncSessionLocal() as db:
        model = TestDatasetJob if run_type == "job" else RemediationReport
        await db.execute(
            update(model)
            .where(model.id == run_id)
            .values(progress_detail=detail)
        )
        await db.commit()


async def _is_run_cancelled(run_type: str, run_id: int) -> bool:
    async with AsyncSessionLocal() as db:
        model = TestDatasetJob if run_type == "job" else RemediationReport
        result = await db.execute(select(model.status).where(model.id == run_id))
        row = result.one_or_none()
        return row is None or row[0] not in ("running", "pending")


def _bundle_doc_prompt(bundle: dict, persona: str, system_name: str) -> tuple[str, str]:
    controls = bundle.get("controls", [])
    intent = bundle.get("document_intent", "implements")
    artifact_type = bundle.get("artifact_type", "procedure")
    evidence_mix = bundle.get("evidence_mix", "balanced")
    evidence_role = bundle.get("evidence_role", "implementation")

    intro = [
        persona,
        f"You are writing a consolidated federal security package component for {system_name}.",
        f"Document title: {bundle.get('title', 'Security Package Component')}",
        f"Artifact type: {artifact_type}",
        f"Document intent: {intent}",
        f"Evidence mix target: {evidence_mix}",
        "",
        "Return ONLY a valid JSON object:",
        '{',
        '  "title": "Full document title",',
        '  "sections": [',
        '    {"type": "heading", "level": 1, "text": "Section heading"},',
        '    {"type": "paragraph", "text": "Plain prose only"},',
        '    {"type": "numbered_list", "items": ["Step"]},',
        '    {"type": "bullet_list", "items": ["Item"]},',
        '    {"type": "table", "headers": ["Col1", "Col2"], "rows": [["v1", "v2"]]}',
        "  ]",
        "}",
        "",
        "Requirements:",
        "- This is one consolidated document covering multiple controls.",
        "- Each control must begin with a 'Control Implementation Summary' heading followed by 2-4 paragraphs of coherent, human-readable control-level narrative before the objective sections.",
        "- Each control must have its own heading and sub-sections for the listed objectives.",
        "- For compliant controls, write current-state implemented language with concrete evidence detail.",
        "- Rewrite objectives as satisfied current-state implementation evidence unless the document intent is explicitly plans.",
        "- Do not repeat negative gap phrasing such as missing, not documented, no evidence, or lacks in the artifact text.",
        "- For passing packages, prioritize useful implementation detail over gap narration.",
        "- Use realistic names, tools, cadences, reviewers, and records.",
        "- No markdown symbols in paragraph text.",
    ]

    if artifact_type == "technical_artifact":
        intro.extend(
            [
                "- This must read like technical implementation evidence, not a policy narrative.",
                "- Name concrete tools, console areas, config settings, schedules, and owners.",
                "- Include tables for configuration values, verification records, or monitored assets where appropriate.",
            ]
        )
    if evidence_role == "validation":
        intro.extend(
            [
                "- Focus on proof and validation: test execution, verification dates, observed results, and reviewer sign-off.",
                "- Include explicit verification methods, sample outputs, and evidence references for each control.",
                "- Treat this as the package's last-mile proof layer for hard technical controls.",
            ]
        )
    elif evidence_role == "architecture":
        intro.extend(
            [
                "- Focus on architecture and system design context: trust boundaries, key components, security tooling, and defense-in-depth.",
                "- Make architecture relationships explicit so the assessor can trace where each control is implemented.",
            ]
        )
    elif evidence_role == "governance":
        intro.extend(
            [
                "- Focus on authority, roles, review cadence, and organizational requirements that govern the implementation.",
            ]
        )
    elif evidence_role == "operations":
        intro.extend(
            [
                "- Focus on how operators execute the control day to day, including tickets, review records, and retained evidence.",
            ]
        )

    control_lines: list[str] = []
    for ctrl in controls:
        contracts = build_control_closure_guidance(
            control_id=ctrl["control_id"],
            control_title=ctrl["title"],
            gaps=ctrl.get("objectives") or [ctrl.get("statement") or ctrl["title"]],
            system_name=system_name,
            mode="synthetic",
        )["objective_contracts"]
        control_lines.extend(
            [
                "",
                f"CONTROL {ctrl['control_id']} - {ctrl['title']}",
                f"Family: {ctrl['family_title']}",
                f"Target status: {ctrl.get('target_status', 'compliant')}",
                build_control_statement_generation_guidance(
                    ctrl["control_id"],
                    ctrl["title"],
                    ctrl.get("objectives") or [ctrl.get("statement") or ctrl["title"]],
                ),
                "Assessment objectives:",
            ]
        )
        objectives = ctrl.get("objectives") or []
        if objectives:
            for idx, objective in enumerate(objectives, start=1):
                control_lines.append(f"  {idx}. {objective}")
        else:
            control_lines.append(f"  1. {ctrl.get('statement', ctrl['title'])[:300]}")
        control_lines.extend(["", format_contracts_for_prompt(contracts)])

    if intent == "plans":
        intro.extend(
            [
                "- This document must read like a POA&M / planned improvements register.",
                "- Use future-state or planned language where appropriate.",
                "- Make clear these controls are intentionally expected to remain unsatisfied.",
            ]
        )

    user_prompt = "\n".join(
        [
            f"Generate the complete document for {system_name}.",
            f"Bundle title: {bundle.get('title', 'Security Package Component')}",
            f"Controls covered: {', '.join(bundle.get('control_ids', []))}",
            *control_lines,
        ]
    )
    return ("\n".join(intro), user_prompt)


def _fallback_bundle_json(bundle: dict, system_name: str) -> dict:
    sections: list[dict] = [
        {"type": "heading", "level": 1, "text": bundle.get("title", "Security Package Component")},
        {
            "type": "paragraph",
            "text": (
                f"This consolidated artifact package supports {system_name} and covers controls "
                f"{', '.join(bundle.get('control_ids', []))}. Each objective below is written as a pass-oriented "
                "closure statement with concrete facts, verification, and retained evidence."
            ),
        },
    ]
    for ctrl in bundle.get("controls", []):
        control_guidance = build_control_closure_guidance(
            control_id=ctrl["control_id"],
            control_title=ctrl["title"],
            gaps=ctrl.get("objectives") or [ctrl.get("statement") or ctrl["title"]],
            system_name=system_name,
            mode="synthetic",
        )
        control_sections = build_contract_sections(
            contracts=control_guidance["objective_contracts"],
            system_name=system_name,
            document_type=bundle.get("document_type") or bundle.get("artifact_type"),
            intro_title=f"{ctrl['control_id']} - {ctrl['title']}",
            intro_text=(
                f"Target assessment status is {ctrl.get('target_status', 'compliant').replace('_', ' ')} for "
                f"{ctrl['control_id']}, and the following sections state what is implemented today."
            ),
        )
        summary_paragraphs = [
            paragraph
            for paragraph in synthesize_control_implementation_statement(
                control_id=ctrl["control_id"],
                control_title=ctrl["title"],
                status=ctrl.get("target_status", "compliant"),
                objectives=ctrl.get("objectives") or [ctrl.get("statement") or ctrl["title"]],
                gap_analysis=[
                    {"objective_id": contract["objective_id"], "met": "yes", "gap": None}
                    for contract in control_guidance["objective_contracts"]
                ],
            ).split("\n\n")
            if paragraph.strip()
        ]
        control_sections[2:2] = [
            {"type": "heading", "level": 2, "text": "Control Implementation Summary"},
            *[
                {"type": "paragraph", "text": paragraph}
                for paragraph in summary_paragraphs
            ],
        ]
        sections.extend(control_sections)
    return {"title": bundle.get("title", "Security Package Component"), "sections": sections}


def _normalize_bundle_sections(bundle: dict, parsed: dict, system_name: str) -> dict:
    sections = parsed.get("sections") if isinstance(parsed, dict) else []
    sections = sections if isinstance(sections, list) else []
    contracts: list[dict] = []
    for ctrl in bundle.get("controls", []):
        contracts.extend(
            build_control_closure_guidance(
                control_id=ctrl["control_id"],
                control_title=ctrl["title"],
                gaps=ctrl.get("objectives") or [ctrl.get("statement") or ctrl["title"]],
                system_name=system_name,
                mode="synthetic",
            )["objective_contracts"]
        )
    contracts_ok, _ = sections_satisfy_contracts(sections, contracts)
    substantive_sections = sum(
        1
        for section in sections
        if section.get("type") in {"heading", "paragraph", "bullet_list", "numbered_list", "table"}
        and any(section.get(key) for key in ("text", "items", "rows"))
    )
    if substantive_sections < 4 or not contracts_ok:
        return _fallback_bundle_json(bundle, system_name)
    return {
        "title": parsed.get("title") or bundle.get("title", "Security Package Component"),
        "sections": sections,
    }


# ── Main entry point ──────────────────────────────────────────────────────────

async def _generate_legacy_test_dataset(job_id: int, project_id: int) -> None:
    """
    Preserved legacy generator, superseded by ``generate_test_dataset`` below.

    Background task: generate a complete ATO test dataset for every control
    in the project's baseline. Deletes docs from any previous job first for
    a clean, self-consistent package.
    """
    from app.services.controls.catalog import load_baseline
    from app.services.llm.runtime import build_provider_for_purpose
    from app.core.config import get_settings

    settings = get_settings()

    # ── Load project + job metadata ──────────────────────────────────────────
    async with AsyncSessionLocal() as db:
        proj_result = await db.execute(select(Project).where(Project.id == project_id))
        project = proj_result.scalar_one_or_none()
        if not project:
            return

        profile_result = await db.execute(
            select(SystemProfile).where(SystemProfile.project_id == project_id)
        )
        profile = profile_result.scalar_one_or_none()

        job_result = await db.execute(
            select(TestDatasetJob).where(TestDatasetJob.id == job_id)
        )
        job = job_result.scalar_one_or_none()
        if not job:
            return

        created_by = job.created_by
        prior_content = job.content_json or {}
        prior_entries: list[dict] = prior_content.get("artifacts", [])
        prior_doc_ids: list[int] = list(job.generated_doc_ids or [])

        # Clean slate: delete docs from the most recent PREVIOUS completed job
        prev_result = await db.execute(
            select(TestDatasetJob)
            .where(
                TestDatasetJob.project_id == project_id,
                TestDatasetJob.id != job_id,
                TestDatasetJob.status.in_(["complete", "failed", "cancelled"]),
            )
            .order_by(TestDatasetJob.created_at.desc())
            .limit(1)
        )
        prev_job = prev_result.scalar_one_or_none()

        # Mark job running
        await db.execute(
            update(TestDatasetJob)
            .where(TestDatasetJob.id == job_id)
            .values(status="running")
        )
        await db.commit()

    # Delete previous job's documents (clean slate)
    if prev_job and prev_job.generated_doc_ids:
        async with AsyncSessionLocal() as db:
            old_docs = await db.execute(
                select(Document).where(Document.id.in_(prev_job.generated_doc_ids))
            )
            for doc in old_docs.scalars().all():
                try:
                    Path(doc.file_path).unlink(missing_ok=True)
                except Exception:
                    pass
                await db.delete(doc)
            await db.commit()
        logger.info(
            "Deleted %d docs from previous job %d for project %d",
            len(prev_job.generated_doc_ids), prev_job.id, project_id,
        )

    total_started_at = perf_counter()
    timings: dict[str, float | int] = {"checkpoint_interval": PARTIAL_SAVE_EVERY_BUNDLES}

    impact_level = project.impact_baseline or "moderate"
    system_name = project.name
    controls = load_baseline(impact_level)
    async with AsyncSessionLocal() as cfg_db:
        llm, _ = await build_provider_for_purpose(
            cfg_db,
            "test_dataset_generation",
            provider_name=getattr(profile, "llm_provider", None) or "ollama",
            model=getattr(profile, "llm_model", None) or None,
        )

    # ── Extract context from existing documents ──────────────────────────────
    await _set_progress(job_id, "Extracting system context from existing documents…")
    ctx = await _extract_context_from_docs(project_id, llm, impact_level)
    # Use real system name if extracted, otherwise fall back to project name
    if ctx.get("system_name"):
        system_name = ctx["system_name"]
    persona = _build_persona(ctx, system_name)

    upload_dir = Path(settings.upload_dir) / str(project_id)
    upload_dir.mkdir(parents=True, exist_ok=True)

    # Resume support — skip controls already processed in THIS job (restart recovery)
    completed: set[str] = {
        e["control_id"] for e in prior_entries
        if "control_id" in e and not e.get("error")
    }
    all_entries: list[dict] = list(prior_entries)
    doc_ids: list[int] = list(prior_doc_ids)
    system_prompt_cache: dict[str, str] = {}

    total = len(controls)
    done = len(completed)

    if completed:
        logger.info("Resuming job %d — %d/%d controls already done", job_id, done, total)

    for ctrl in controls:
        cid = ctrl.display_id

        if await _is_cancelled(job_id):
            logger.info("Job %d cancelled at control %s", job_id, cid)
            return

        if cid in completed:
            continue

        done += 1
        await _set_progress(job_id, f"▶ {done}/{total} — {cid}: {ctrl.title[:55]}…")

        doc_type_label, doc_type_field, content_mode = _classify_control(cid)
        objectives = _objectives_from_control(ctrl)
        objectives_text = _format_objectives(objectives, cid, ctrl.title)
        statement_guidance = build_control_statement_generation_guidance(cid, ctrl.title, objectives)
        summary_paragraphs = [
            paragraph
            for paragraph in synthesize_control_implementation_statement(
                control_id=cid,
                control_title=ctrl.title,
                status="compliant",
                objectives=objectives,
                gap_analysis=[
                    {"objective_id": obj["objective_id"], "met": "yes", "gap": None}
                    for obj in objectives
                ],
            ).split("\n\n")
            if paragraph.strip()
        ]
        system_prompt = system_prompt_cache.setdefault(
            content_mode,
            _build_prompt(content_mode, persona, system_name),
        )

        user_msg = (
            f"{objectives_text}\n\n"
            f"{statement_guidance}\n\n"
            f"Write a complete {doc_type_label} for {system_name} that satisfies every "
            f"objective listed above. Each objective MUST get its own labeled section."
        )

        content_json_str = "{}"
        try:
            raw = await _complete_with_timeout(
                llm,
                system_prompt,
                user_msg,
                timeout_secs=BUNDLE_GENERATION_TIMEOUT_SECS,
                label=f"Control generation {cid}",
            )
            s = raw.find("{")
            e = raw.rfind("}") + 1
            if s != -1 and e > s:
                content_json_str = raw[s:e]
            else:
                raise ValueError("No JSON in LLM response")
        except Exception as ex:
            logger.warning("LLM failed for %s: %s", cid, ex)
            summary_paragraphs = [
                paragraph
                for paragraph in synthesize_control_implementation_statement(
                    control_id=cid,
                    control_title=ctrl.title,
                    status="compliant",
                    objectives=objectives,
                    gap_analysis=[
                        {"objective_id": obj["objective_id"], "met": "yes", "gap": None}
                        for obj in objectives
                    ],
                ).split("\n\n")
                if paragraph.strip()
            ]
            # Minimal deterministic fallback — still produces an indexable document
            content_json_str = json.dumps({
                "title": f"{cid} — {ctrl.title} — {doc_type_label}",
                "sections": [
                    {"type": "paragraph", "text": (
                        f"This document provides the {doc_type_label} for {system_name} "
                        f"addressing NIST SP 800-53 Rev 5 control {cid} — {ctrl.title}."
                    )},
                    {"type": "heading", "level": 2, "text": "Control Implementation Summary"},
                    *[
                        {"type": "paragraph", "text": paragraph}
                        for paragraph in summary_paragraphs
                    ],
                ] + [
                    {
                        "type": "paragraph",
                        "text": (
                            f"This section satisfies NIST 800-53A assessment objective "
                            f"{obj['objective_id']}. "
                            f"{system_name} implements the following: {obj['description']}"
                        ),
                    }
                    for obj in objectives
                ],
            })

        try:
            parsed = json.loads(content_json_str)
        except json.JSONDecodeError:
            parsed = {"title": f"{cid} {doc_type_label}", "sections": []}

        doc_title = parsed.get("title") or f"{cid} — {ctrl.title} — {doc_type_label}"
        sections = parsed.get("sections") or []
        has_summary_heading = any(
            section.get("type") == "heading"
            and str(section.get("text") or "").strip().lower() == "control implementation summary"
            for section in sections
            if isinstance(section, dict)
        )
        if not has_summary_heading:
            sections = [
                {"type": "heading", "level": 2, "text": "Control Implementation Summary"},
                *[
                    {"type": "paragraph", "text": paragraph}
                    for paragraph in summary_paragraphs
                ],
                *sections,
            ]
        org = ctx.get("organization", "Keystone Federal Solutions, LLC")
        doc_bytes = _build_docx(doc_title, sections, org)

        filename = _build_generated_filename(
            "TESTDATA",
            f"{cid}_{doc_type_field}",
            f"{ctrl.title}_{doc_type_label}",
        )

        entry: dict = {
            "control_id": cid,
            "family": ctrl.family_id.upper(),
            "title": doc_title,
            "artifact_type": doc_type_field,
        }

        try:
            doc_id = await _save_doc(
                file_bytes=doc_bytes,
                filename=filename,
                project_id=project_id,
                upload_dir=upload_dir,
                created_by=created_by,
                control_id=cid,
                document_type=doc_type_field,
            )
            doc_ids.append(doc_id)
            entry["doc_id"] = doc_id
        except Exception as save_err:
            logger.error("Save failed for %s: %s", cid, save_err)
            entry["error"] = str(save_err)

        all_entries.append(entry)
        completed.add(cid)

        # Incremental save every control for restart recovery
        _partial = {
            "report_type": "test_dataset",
            "generated_at": datetime.now(UTC).isoformat(),
            "system_name": system_name,
            "summary": {
                "controls_addressed": done,
                "total_controls": total,
                "documents_created": len(doc_ids),
                "impact_baseline": impact_level.upper(),
            },
            "artifacts": all_entries,
        }
        try:
            async with AsyncSessionLocal() as _db:
                await _db.execute(
                    update(TestDatasetJob)
                    .where(TestDatasetJob.id == job_id)
                    .values(content_json=_partial, generated_doc_ids=doc_ids)
                )
                await _db.commit()
        except Exception:
            pass  # non-fatal

    # ── Final completion ─────────────────────────────────────────────────────
    final_content = {
        "report_type": "test_dataset",
        "generated_at": datetime.now(UTC).isoformat(),
        "system_name": system_name,
        "organization": ctx.get("organization", ""),
        "summary": {
            "controls_addressed": total,
            "documents_created": len(doc_ids),
            "impact_baseline": impact_level.upper(),
            "context_source": "extracted" if ctx.get("system_name") else "default_persona",
        },
        "artifacts": all_entries,
    }

    async with AsyncSessionLocal() as db:
        await db.execute(
            update(TestDatasetJob)
            .where(TestDatasetJob.id == job_id)
            .values(
                status="complete",
                content_json=final_content,
                generated_doc_ids=doc_ids,
                progress_detail=(
                    f"Complete — {len(doc_ids)} documents generated for {total} controls "
                    f"({impact_level.upper()} baseline), indexed and ready for assessment"
                ),
            )
        )
        await db.commit()

    logger.info(
        "Test dataset job %d complete: %d docs for %d controls (%s)",
        job_id, len(doc_ids), total, impact_level.upper(),
    )


async def generate_test_dataset(run_id: int, project_id: int | None = None) -> None:
    """Generate a consolidated fictitious ATO package for either a job or remediation report."""
    from app.core.config import get_settings
    from app.services.controls.catalog import load_baseline
    from app.services.llm.runtime import build_provider_for_purpose
    from app.services.artifact_validation import validate_generated_artifacts
    from app.services.evidence_quality import (
        enhance_artifact_document,
        evidence_repair_prompt,
    )
    from app.services.system_knowledge import extract_system_knowledge

    settings = get_settings()
    run_type = "job" if project_id is not None else "report"

    async with AsyncSessionLocal() as db:
        if run_type == "job":
            project = (await db.execute(select(Project).where(Project.id == project_id))).scalar_one_or_none()
            profile = (await db.execute(select(SystemProfile).where(SystemProfile.project_id == project_id))).scalar_one_or_none()
            run_record = (await db.execute(select(TestDatasetJob).where(TestDatasetJob.id == run_id))).scalar_one_or_none()
            source_assessment_id = None
            source_report_id = None
            previous_run = (
                await db.execute(
                    select(TestDatasetJob)
                    .where(
                        TestDatasetJob.project_id == project_id,
                        TestDatasetJob.id != run_id,
                        TestDatasetJob.status.in_(["complete", "failed", "cancelled"]),
                    )
                    .order_by(TestDatasetJob.created_at.desc())
                    .limit(1)
                )
            ).scalar_one_or_none()
            update_model = TestDatasetJob
        else:
            run_record = (await db.execute(select(RemediationReport).where(RemediationReport.id == run_id))).scalar_one_or_none()
            if run_record is None:
                return
            assessment = (await db.execute(select(Assessment).where(Assessment.id == run_record.assessment_id))).scalar_one_or_none()
            if assessment is None:
                return
            project_id = assessment.project_id
            project = (await db.execute(select(Project).where(Project.id == project_id))).scalar_one_or_none()
            profile = (await db.execute(select(SystemProfile).where(SystemProfile.project_id == project_id))).scalar_one_or_none()
            source_assessment_id = assessment.id
            source_report_id = run_record.id
            previous_run = (
                await db.execute(
                    select(RemediationReport)
                    .where(
                        RemediationReport.assessment_id == assessment.id,
                        RemediationReport.report_type == "test_dataset",
                        RemediationReport.id != run_id,
                        RemediationReport.status.in_(["complete", "failed", "cancelled"]),
                    )
                    .order_by(RemediationReport.created_at.desc())
                    .limit(1)
                )
            ).scalar_one_or_none()
            update_model = RemediationReport

        if not project or not run_record:
            return

        created_by = run_record.created_by
        prior_content = run_record.content_json or {}
        config = prior_content.get("config", {})
        prior_entries: list[dict] = list(prior_content.get("artifacts", []))
        prior_doc_ids: list[int] = list(run_record.generated_doc_ids or [])

        await db.execute(update(update_model).where(update_model.id == run_id).values(status="running"))
        await db.commit()

    if previous_run and previous_run.generated_doc_ids:
        async with AsyncSessionLocal() as db:
            old_docs = await db.execute(select(Document).where(Document.id.in_(previous_run.generated_doc_ids)))
            for doc in old_docs.scalars().all():
                try:
                    Path(doc.file_path).unlink(missing_ok=True)
                except Exception:
                    pass
                await db.delete(doc)
            await db.commit()

    impact_level = project.impact_baseline or "moderate"
    system_name = project.name
    controls = load_baseline(impact_level)
    async with AsyncSessionLocal() as cfg_db:
        llm, _ = await build_provider_for_purpose(
            cfg_db,
            "test_dataset_generation",
            provider_name=getattr(profile, "llm_provider", None) or "ollama",
            model=getattr(profile, "llm_model", None) or None,
        )

    total_started_at = perf_counter()
    timings: dict[str, float | int] = {"checkpoint_interval": PARTIAL_SAVE_EVERY_BUNDLES}
    planning_started_at = perf_counter()
    package_style = (config.get("package_style") or "standard").lower()
    evidence_mix = (config.get("evidence_mix") or "balanced").lower()
    target_profile = (config.get("target_profile") or "passing_ato").lower()
    expected_outcomes = build_expected_outcomes(
        [ctrl.display_id for ctrl in controls],
        target_profile=target_profile,
        expected_satisfied_pct=config.get("expected_satisfied_pct"),
        expected_partial_pct=config.get("expected_partial_pct"),
        expected_failed_pct=config.get("expected_failed_pct"),
        family_overrides=config.get("family_overrides"),
        control_overrides=config.get("control_overrides"),
    )
    blueprint = plan_test_dataset_bundles(
        controls,
        package_style=package_style,
        evidence_mix=evidence_mix,
        expected_outcomes=expected_outcomes,
    )
    validation = build_blueprint_validation(blueprint, expected_outcomes)
    timings["planning_secs"] = round(perf_counter() - planning_started_at, 3)

    await _set_run_progress(run_type, run_id, "Planning consolidated test package...")
    context_started_at = perf_counter()
    ctx = await _extract_context_from_docs(project_id, llm, impact_level)
    if ctx.get("system_name"):
        system_name = ctx["system_name"]
    persona = _build_persona(ctx, system_name)
    timings["context_extraction_secs"] = round(perf_counter() - context_started_at, 3)

    upload_dir = Path(settings.upload_dir) / str(project_id)
    upload_dir.mkdir(parents=True, exist_ok=True)

    bundles = blueprint.get("bundles", [])
    completed_bundles = {
        entry.get("bundle_id")
        for entry in prior_entries
        if entry.get("bundle_id") and not entry.get("error")
    }
    artifacts: list[dict] = list(prior_entries)
    doc_ids: list[int] = list(prior_doc_ids)
    pending_dispatch_doc_ids: list[int] = []
    generation_started_at = perf_counter()
    processed_since_checkpoint = 0
    bundle_fallback_cache: dict[str, dict] = {}

    for idx, bundle in enumerate(bundles, start=1):
        if await _is_run_cancelled(run_type, run_id):
            return
        if bundle["bundle_id"] in completed_bundles:
            continue

        await _set_run_progress(run_type, run_id, f"{idx}/{len(bundles)} - generating {bundle['title']}...")

        system_prompt, user_prompt = _bundle_doc_prompt(bundle, persona, system_name)
        fallback_parsed = bundle_fallback_cache.setdefault(
            bundle["bundle_id"],
            _fallback_bundle_json(bundle, system_name),
        )
        parsed = fallback_parsed
        try:
            raw = await _complete_with_timeout(
                llm,
                system_prompt,
                user_prompt,
                timeout_secs=BUNDLE_GENERATION_TIMEOUT_SECS,
                label=f"Bundle generation {bundle['bundle_id']}",
            )
            start = raw.find("{")
            end = raw.rfind("}") + 1
            if start != -1 and end > start:
                parsed = json.loads(raw[start:end])
        except Exception as exc:
            logger.warning("Bundle generation failed for %s: %s", bundle["bundle_id"], exc)

        parsed = _normalize_bundle_sections(bundle, parsed, system_name)
        quality_controls = [
            {
                "control_id": ctrl.get("control_id"),
                "title": ctrl.get("title"),
                "family": ctrl.get("family") or bundle.get("family"),
                "target_status": ctrl.get("target_status", "compliant"),
                "objectives": ctrl.get("objectives") or [ctrl.get("statement") or ctrl.get("title")],
            }
            for ctrl in bundle.get("controls", [])
        ]
        parsed, evidence_quality = enhance_artifact_document(
            document=parsed,
            controls=quality_controls,
            artifact_type=bundle.get("artifact_type", "procedure"),
            evidence_role=bundle.get("evidence_role", "implementation"),
            system_name=system_name,
            system_context=persona,
            organization=ctx.get("organization") or FALLBACK_PERSONA["organization"],
            source="synthetic_test",
        )
        if not evidence_quality.get("passed"):
            try:
                repair_system, repair_user = evidence_repair_prompt(
                    document=parsed,
                    controls=quality_controls,
                    quality_summary=evidence_quality,
                    system_name=system_name,
                )
                repair_raw = await _complete_with_timeout(
                    llm,
                    repair_system,
                    repair_user,
                    timeout_secs=BUNDLE_REPAIR_TIMEOUT_SECS,
                    label=f"Bundle repair {bundle['bundle_id']}",
                )
                start = repair_raw.find("{")
                end = repair_raw.rfind("}") + 1
                if start != -1 and end > start:
                    repaired = json.loads(repair_raw[start:end])
                    parsed, evidence_quality = enhance_artifact_document(
                        document=_normalize_bundle_sections(bundle, repaired, system_name),
                        controls=quality_controls,
                        artifact_type=bundle.get("artifact_type", "procedure"),
                        evidence_role=bundle.get("evidence_role", "implementation"),
                        system_name=system_name,
                        system_context=persona,
                        organization=ctx.get("organization") or FALLBACK_PERSONA["organization"],
                        source="synthetic_test_repaired",
                    )
            except Exception as exc:
                logger.warning("Evidence quality repair failed for %s: %s", bundle["bundle_id"], exc)
        doc_title = parsed.get("title") or bundle["title"]
        sections = parsed.get("sections") or fallback_parsed["sections"]
        doc_bytes = _build_docx(doc_title, sections, ctx.get("organization", FALLBACK_PERSONA["organization"]))
        filename = _build_generated_filename(
            "TESTPKG",
            bundle["bundle_id"],
            bundle["title"],
        )

        artifact_entry = {
            "bundle_id": bundle["bundle_id"],
            "family": bundle["family"],
            "title": doc_title,
            "artifact_type": bundle["artifact_type"],
            "document_type": bundle["document_type"],
            "document_intent": bundle["document_intent"],
            "evidence_role": bundle.get("evidence_role", "implementation"),
            "controls_addressed": bundle["control_ids"],
            "target_statuses": {
                ctrl["control_id"]: ctrl.get("target_status", "compliant") for ctrl in bundle.get("controls", [])
            },
            "evidence_quality": evidence_quality,
        }

        try:
            doc_id = await _save_doc(
                file_bytes=doc_bytes,
                filename=filename,
                project_id=project_id,
                upload_dir=upload_dir,
                created_by=created_by,
                control_id=bundle["control_ids"][0],
                document_type=bundle["document_type"],
                document_intent=bundle["document_intent"],
                controls_addressed=bundle["control_ids"],
                source_assessment_id=source_assessment_id,
                source_remediation_report_id=source_report_id,
                trigger_parse=False,
            )
            artifact_entry["doc_id"] = doc_id
            doc_ids.append(doc_id)
            pending_dispatch_doc_ids.append(doc_id)
        except Exception as save_err:
            artifact_entry["error"] = str(save_err)
            logger.error("Save failed for bundle %s: %s", bundle["bundle_id"], save_err)

        artifacts.append(artifact_entry)
        completed_bundles.add(bundle["bundle_id"])

        processed_since_checkpoint += 1
        if processed_since_checkpoint >= PARTIAL_SAVE_EVERY_BUNDLES:
            partial_content = {
                "report_type": "test_dataset",
                "config": config,
                "generated_at": datetime.now(UTC).isoformat(),
                "system_name": system_name,
                "organization": ctx.get("organization", ""),
                "expected_outcomes": expected_outcomes,
                "blueprint": blueprint,
                "validation": validation,
                "summary": {
                    "controls_addressed": len(expected_outcomes.get("by_control", {})),
                    "total_controls": len(controls),
                    "documents_created": len(doc_ids),
                    "impact_baseline": impact_level.upper(),
                    "package_style": package_style,
                    "target_profile": target_profile,
                },
                "artifacts": artifacts,
            }
            async with AsyncSessionLocal() as db:
                await db.execute(
                    update(update_model)
                    .where(update_model.id == run_id)
                    .values(content_json=partial_content, generated_doc_ids=doc_ids)
                )
                await db.commit()
            if pending_dispatch_doc_ids:
                dispatch_started_at = perf_counter()
                from app.services.parsers.dispatcher import dispatch_parse_batch
                await dispatch_parse_batch(pending_dispatch_doc_ids)
                timings["index_dispatch_secs"] = round(
                    (timings.get("index_dispatch_secs", 0.0) or 0.0)
                    + (perf_counter() - dispatch_started_at),
                    3,
                )
                pending_dispatch_doc_ids = []
            processed_since_checkpoint = 0

    timings["generation_secs"] = round(perf_counter() - generation_started_at, 3)

    if pending_dispatch_doc_ids:
        dispatch_started_at = perf_counter()
        from app.services.parsers.dispatcher import dispatch_parse_batch
        await dispatch_parse_batch(pending_dispatch_doc_ids)
        timings["index_dispatch_secs"] = round(
            (timings.get("index_dispatch_secs", 0.0) or 0.0)
            + (perf_counter() - dispatch_started_at),
            3,
        )
    elif "index_dispatch_secs" not in timings:
        timings["index_dispatch_secs"] = 0.0

    validation_summary = None
    system_knowledge_summary = None
    if doc_ids:
        async with AsyncSessionLocal() as db:
            validation_started_at = perf_counter()
            validation_summary = await validate_generated_artifacts(
                db,
                project_id=project_id,
                document_ids=doc_ids,
                source_mode="synthetic_test",
                source_run_id=run_id,
                expected_profile=target_profile,
            )
            timings["artifact_validation_secs"] = round(perf_counter() - validation_started_at, 3)
            knowledge_started_at = perf_counter()
            system_knowledge_summary = await extract_system_knowledge(
                db,
                project_id=project_id,
                source_mode="synthetic_test",
                source_run_id=run_id,
                document_ids=doc_ids,
            )
            timings["system_knowledge_secs"] = round(perf_counter() - knowledge_started_at, 3)
    else:
        timings["artifact_validation_secs"] = 0.0
        timings["system_knowledge_secs"] = 0.0

    timings["total_secs"] = round(perf_counter() - total_started_at, 3)

    final_content = {
        "report_type": "test_dataset",
        "generation_mode": "synthetic_test",
        "config": config,
        "generated_at": datetime.now(UTC).isoformat(),
        "system_name": system_name,
        "organization": ctx.get("organization", ""),
        "expected_outcomes": expected_outcomes,
        "blueprint": blueprint,
        "validation": validation,
        "artifact_validation": validation_summary,
        "system_knowledge": system_knowledge_summary,
        "timing": timings,
        "summary": {
            "controls_addressed": len(expected_outcomes.get("by_control", {})),
            "total_controls": len(controls),
            "documents_created": len(doc_ids),
            "impact_baseline": impact_level.upper(),
            "context_source": "extracted" if ctx.get("system_name") else "default_persona",
            "package_style": package_style,
            "target_profile": target_profile,
        },
        "artifacts": artifacts,
    }

    async with AsyncSessionLocal() as db:
        await db.execute(
            update(update_model)
            .where(update_model.id == run_id)
            .values(
                status="complete",
                content_json=final_content,
                generated_doc_ids=doc_ids,
                progress_detail=(
                    f"Complete - {len(doc_ids)} consolidated package documents generated "
                    f"for {len(controls)} controls ({impact_level.upper()} baseline)"
                ),
            )
        )
        await db.commit()

    logger.info(
        "Test dataset %s %d complete: %d docs for %d controls (%s)",
        run_type,
        run_id,
        len(doc_ids),
        len(controls),
        impact_level.upper(),
    )
