"""Word (.docx) parser using python-docx."""
from __future__ import annotations

from pathlib import Path

from app.services.parsers.base import ParsedBlock, ParsedCell, ParsedDocument, ParsedPage


def parse_docx(file_path: str) -> ParsedDocument:
    try:
        from docx import Document
        from docx.document import Document as DocumentType
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        return ParsedDocument(
            filename=Path(file_path).name,
            file_type="docx",
            error="python-docx not installed",
        )

    def iter_block_items(parent):
        parent_elm = parent.element.body if isinstance(parent, DocumentType) else parent._tc
        for child in parent_elm.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, parent)
            elif child.tag.endswith("}tbl"):
                yield Table(child, parent)

    def normalize_heading_path(path: list[str], heading: str, style_name: str) -> list[str]:
        if not heading:
            return path
        if "heading" in style_name:
            parts = style_name.split()
            level = None
            for part in reversed(parts):
                if part.isdigit():
                    level = int(part)
                    break
            if level:
                return path[: level - 1] + [heading]
        return path + [heading] if not path or path[-1] != heading else path

    doc = ParsedDocument(
        filename=Path(file_path).name,
        file_type="docx",
        parser_name="python-docx",
        parser_version="ordered-blocks-v1",
    )
    try:
        document = Document(file_path)
        doc.metadata = {
            "title": document.core_properties.title or "",
            "author": document.core_properties.author or "",
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
        }

        page = ParsedPage(page_number=1, content="")
        heading_path: list[str] = []
        content_lines: list[str] = []
        block_counter = 0
        table_counter = 0

        for item in iter_block_items(document):
            if isinstance(item, Paragraph):
                text = item.text.strip()
                if not text:
                    continue
                style_name = item.style.name.lower() if item.style else ""
                block_type = "heading" if "heading" in style_name else "list_item" if style_name.startswith("list") else "paragraph"
                if block_type == "heading":
                    heading_path = normalize_heading_path(heading_path, text, style_name)
                    page.section_title = page.section_title or text
                block_counter += 1
                page.blocks.append(
                    ParsedBlock(
                        block_id=f"docx-b{block_counter}",
                        block_type=block_type,
                        text=text,
                        heading_path=list(heading_path),
                        metadata={"style_name": item.style.name if item.style else None},
                    )
                )
                content_lines.append(text)
                continue

            table_counter += 1
            table_id = f"docx-table-{table_counter}"
            header_cells: list[str | None] = []
            for row_idx, row in enumerate(item.rows):
                row_cells: list[ParsedCell] = []
                row_text_parts: list[str] = []
                for col_idx, cell in enumerate(row.cells):
                    cell_text = "\n".join(
                        p.text.strip() for p in cell.paragraphs if p.text.strip()
                    ).strip()
                    if row_idx == 0:
                        header = cell_text or None
                        header_cells.append(header)
                    else:
                        header = header_cells[col_idx] if col_idx < len(header_cells) else None
                    row_cells.append(
                        ParsedCell(
                            text=cell_text,
                            row_index=row_idx,
                            col_index=col_idx,
                            header=header,
                        )
                    )
                    if cell_text:
                        row_text_parts.append(f"{header}: {cell_text}" if header and row_idx > 0 else cell_text)

                block_counter += 1
                row_text = " | ".join(part for part in row_text_parts if part)
                if row_text:
                    page.blocks.append(
                        ParsedBlock(
                            block_id=f"{table_id}-r{row_idx}",
                            block_type="table_row",
                            text=row_text,
                            heading_path=list(heading_path),
                            row_index=row_idx,
                            table_id=table_id,
                            cells=row_cells,
                            metadata={"column_count": len(row.cells)},
                        )
                    )
                    content_lines.append(row_text)

        page.content = "\n".join(content_lines)
        doc.pages = [page]
        doc.full_text = page.content
    except Exception as e:
        doc.error = str(e)
    return doc
