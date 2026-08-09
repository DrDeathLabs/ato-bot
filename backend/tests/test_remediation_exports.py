import io
import zipfile

from docx import Document

from app.services.remediation_service import _build_artifact_package_docx, _build_guide_docx


def _docx_text(payload: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        assert package.testzip() is None
    document = Document(io.BytesIO(payload))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        parts.extend(cell.text for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def test_remediation_guide_docx_uses_current_structured_sections():
    payload = _build_guide_docx(
        [
            {
                "family": "AC",
                "family_title": "Access Control",
                "actions": [
                    {
                        "control_id": "AC-2",
                        "gap": "Account review evidence is missing",
                        "action": "Export the current account review record.",
                        "responsible": "IAM Lead",
                        "effort": "One week",
                        "success_criteria": "A dated review record is retained.",
                    }
                ],
            }
        ],
        {
            "non_compliant": 1,
            "partially_compliant": 0,
            "families_affected": ["AC"],
            "total_actions": 1,
        },
        "E2E System",
        "August 9, 2026",
    )
    text = _docx_text(payload)
    assert "AC — Access Control" in text
    assert "Export the current account review record" in text


def test_artifact_package_docx_builds_without_legacy_markdown_helpers():
    payload = _build_artifact_package_docx(
        {
            "system_name": "E2E System",
            "artifacts": [
                {
                    "family": "AC",
                    "title": "Access Control Remediation Package",
                    "controls_addressed": ["AC-2", "AC-6"],
                }
            ],
        },
        "August 9, 2026",
    )
    text = _docx_text(payload)
    assert "Access Control Remediation Package" in text
    assert "AC-2, AC-6" in text
